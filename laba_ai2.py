# app.py
import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from scipy.integrate import solve_ivp
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from io import BytesIO
import base64

# =========================
# Данные вариантов
# =========================
variants_data = {
    1: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
        "x0": 0.2, "y0": -0.8},
    2: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
        "x0": 0.6, "y0": 0.6},
    3: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
        "x0": -0.4, "y0": 0.8},
    4: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
        "x0": -1.2, "y0": -1.8},
    5: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
        "x0": -0.7, "y0": -0.4},
    6: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
        "x0": -1.6, "y0": 0.6},
    7: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
        "x0": 1.2, "y0": 0.4},
    8: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
        "x0": 2.6, "y0": -1.2},
    9: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
        "x0": 1.4, "y0": -1.8},
    10: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 3.2, "y0": 0.3},
    11: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -1.7, "y0": 0.6},
    12: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -3.6, "y0": -0.8},
    13: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -0.2, "y0": -1.2},
    14: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -0.6, "y0": 0.8},
    15: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 0.4, "y0": 0.7},
    16: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 1.2, "y0": -0.4},
    17: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 0.7, "y0": -1.2},
    18: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 1.6, "y0": 0.4},
    19: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -1.2, "y0": 0.6},
    20: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -2.6, "y0": -1.6},
    21: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -1.4, "y0": -0.4},
    22: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -3.2, "y0": 0.8},
    23: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 1.7, "y0": 0.2},
    24: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 3.6, "y0": -1.2},
    25: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 0.3, "y0": -1.8},
    26: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 0.8, "y0": 0.2},
    27: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -0.6, "y0": 0.8},
    28: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -1.4, "y0": -1.8},
    29: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -0.8, "y0": -1.2},
    30: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": -2.4, "y0": 0.4},
    31: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 1.3, "y0": 0.4},
    32: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-4, -2, 2, 4], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 2.8, "y0": -1.6},
    33: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-4, -2, 2, 4], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": 1.6, "y0": -0.4},
    34: {"I1": [-6, -4, -2, 2, 4, 6], "I2": [-2, -1, 1, 2], "O": [-3, -2, -1, 1, 2, 3], "C": [-2, -1, 0, 1, 2],
         "x0": 3.4, "y0": 0.4},
    35: {"I1": [-3, -2, -1, 1, 2, 3], "I2": [-2, -1, 1, 2], "O": [-6, -4, -2, 2, 4, 6], "C": [-4, -2, 0, 2, 4],
         "x0": -1.8, "y0": 0.6},
}

# База правил (из таблицы)
rules = [
    ("I2P", "I1NH", "OZ"), ("I2P", "I1NL", "OZ"), ("I2P", "I1Z", "OPL"),
    ("I2P", "I1PL", "OPH"), ("I2P", "I1PH", "OPH"),
    ("I2Z", "I1NH", "ONL"), ("I2Z", "I1NL", "ONL"), ("I2Z", "I1Z", "OZ"),
    ("I2Z", "I1PL", "OPL"), ("I2Z", "I1PH", "OPL"),
    ("I2N", "I1NH", "ONH"), ("I2N", "I1NL", "ONH"), ("I2N", "I1Z", "ONL"),
    ("I2N", "I1PL", "OZ"), ("I2N", "I1PH", "OZ")
]


# =========================
# Функции принадлежности
# =========================
def z_mf(x, a, b):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    y[x <= a] = 1.0
    y[x >= b] = 0.0
    mask = (x > a) & (x < b)
    y[mask] = (b - x[mask]) / (b - a)
    return np.clip(y, 0.0, 1.0)


def s_mf(x, a, b):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    y[x <= a] = 0.0
    y[x >= b] = 1.0
    mask = (x > a) & (x < b)
    y[mask] = (x[mask] - a) / (b - a)
    return np.clip(y, 0.0, 1.0)


def tri_mf(x, a, b, c):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    left = (x >= a) & (x <= b)
    right = (x >= b) & (x <= c)
    if b != a:
        y[left] = (x[left] - a) / (b - a)
    if c != b:
        y[right] = np.maximum(y[right], (c - x[right]) / (c - b))
    y[x == b] = 1.0
    return np.clip(y, 0.0, 1.0)


def i1_sets(x, params):
    I1_min, I1_1, I1_2, I1_3, I1_4, I1_max = params
    return {
        "I1NH": z_mf(x, I1_1, I1_2),
        "I1NL": tri_mf(x, I1_1, I1_2, 0),
        "I1Z": tri_mf(x, I1_2, 0, I1_3),
        "I1PL": tri_mf(x, 0, I1_3, I1_4),
        "I1PH": s_mf(x, I1_3, I1_4)
    }


def i2_sets(x, params):
    I2_min, I2_1, I2_2, I2_max = params
    return {
        "I2N": z_mf(x, I2_1, 0),
        "I2Z": tri_mf(x, I2_1, 0, I2_2),
        "I2P": s_mf(x, 0, I2_2)
    }


def o_sets(x, params):
    O_min, O_1, O_2, O_3, O_4, O_max = params
    return {
        "ONH": z_mf(x, O_1, O_2),
        "ONL": tri_mf(x, O_1, O_2, 0),
        "OZ": tri_mf(x, O_2, 0, O_3),
        "OPL": tri_mf(x, 0, O_3, O_4),
        "OPH": s_mf(x, O_3, O_4)
    }


# =========================
# Функция генерации кода для варианта
# =========================
def generate_code_for_variant(variant, data, C):
    """Генерирует Python код программы для выбранного варианта"""

    i1_params = data["I1"]
    i2_params = data["I2"]
    o_params = data["O"]
    x0, y0 = data["x0"], data["y0"]

    code = f'''import numpy as np
import matplotlib.pyplot as plt
from scipy.integrate import solve_ivp

I1_min, I1_1, I1_2, I1_3, I1_4, I1_max = {i1_params[0]}, {i1_params[1]}, {i1_params[2]}, {i1_params[3]}, {i1_params[4]}, {i1_params[5]}
I2_min, I2_1, I2_2, I2_max = {i2_params[0]}, {i2_params[1]}, {i2_params[2]}, {i2_params[3]}
O_min, O_1, O_2, O_3, O_4, O_max = {o_params[0]}, {o_params[1]}, {o_params[2]}, {o_params[3]}, {o_params[4]}, {o_params[5]}

x0 = {x0}
y0 = {y0}


# Функции принадлежности
def z_mf(x, a, b):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    y[x <= a] = 1.0
    y[x >= b] = 0.0
    mask = (x > a) & (x < b)
    y[mask] = (b - x[mask]) / (b - a)
    return np.clip(y, 0.0, 1.0)


def s_mf(x, a, b):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    y[x <= a] = 0.0
    y[x >= b] = 1.0
    mask = (x > a) & (x < b)
    y[mask] = (x[mask] - a) / (b - a)
    return np.clip(y, 0.0, 1.0)


def tri_mf(x, a, b, c):
    x = np.asarray(x, dtype=float)
    y = np.zeros_like(x)
    left = (x >= a) & (x <= b)
    right = (x >= b) & (x <= c)
    if b != a:
        y[left] = (x[left] - a) / (b - a)
    if c != b:
        y[right] = np.maximum(y[right], (c - x[right]) / (c - b))
    y[x == b] = 1.0
    return np.clip(y, 0.0, 1.0)


def i1_sets(x):
    return {{
        "I1NH": z_mf(x, I1_1, I1_2),
        "I1NL": tri_mf(x, I1_1, I1_2, 0),
        "I1Z": tri_mf(x, I1_2, 0, I1_3),
        "I1PL": tri_mf(x, 0, I1_3, I1_4),
        "I1PH": s_mf(x, I1_3, I1_4)
    }}


def i2_sets(x):
    return {{
        "I2N": z_mf(x, I2_1, 0),
        "I2Z": tri_mf(x, I2_1, 0, I2_2),
        "I2P": s_mf(x, 0, I2_2)
    }}


def o_sets(x):
    return {{
        "ONH": z_mf(x, O_1, O_2),
        "ONL": tri_mf(x, O_1, O_2, 0),
        "OZ": tri_mf(x, O_2, 0, O_3),
        "OPL": tri_mf(x, 0, O_3, O_4),
        "OPH": s_mf(x, O_3, O_4)
    }}


C = {{
    "ONH": {C["ONH"]},
    "ONL": {C["ONL"]},
    "OZ": {C["OZ"]},
    "OPL": {C["OPL"]},
    "OPH": {C["OPH"]}
}}

# База правил (из таблицы)
rules = [
    ("I2P", "I1NH", "OZ"), ("I2P", "I1NL", "OZ"), ("I2P", "I1Z", "OPL"),
    ("I2P", "I1PL", "OPH"), ("I2P", "I1PH", "OPH"),
    ("I2Z", "I1NH", "ONL"), ("I2Z", "I1NL", "ONL"), ("I2Z", "I1Z", "OZ"),
    ("I2Z", "I1PL", "OPL"), ("I2Z", "I1PH", "OPL"),
    ("I2N", "I1NH", "ONH"), ("I2N", "I1NL", "ONH"), ("I2N", "I1Z", "ONL"),
    ("I2N", "I1PL", "OZ"), ("I2N", "I1PH", "OZ")
]


def mamdani_control(e, de, scale_e=1.0, scale_de=1.0, resolution=200):
    """
    e, de – ошибка и её производная
    scale_e, scale_de – масштабирование входов
    Возвращает управляющий сигнал (дефаззификация центром тяжести)
    """
    e_scaled = e / scale_e
    de_scaled = de / scale_de

    # Фаззификация
    mu_i1 = {{k: float(v[0]) for k, v in i1_sets(np.array([e_scaled])).items()}}
    mu_i2 = {{k: float(v[0]) for k, v in i2_sets(np.array([de_scaled])).items()}}

    # Дискретизация выходного диапазона
    z = np.linspace(O_min, O_max, resolution)
    out_mfs = o_sets(z)

    aggregated = np.zeros_like(z)

    for i2_term, i1_term, o_term in rules:
        alpha = min(mu_i1[i1_term], mu_i2[i2_term])  # Степень срабатывания правила
        if alpha > 0:
            clipped = np.minimum(alpha, out_mfs[o_term])
            aggregated = np.maximum(aggregated, clipped)

    if np.sum(aggregated) == 0:
        return 0.0
    return np.sum(z * aggregated) / np.sum(aggregated)  # Центр тяжести


def sugeno_zero_order_control(e, de, scale_e=1.0, scale_de=1.0):
    """
    Вывод по методу Сугено 0-го порядка
    Для заданных e и de возвращает взвешенное среднее констант
    """
    e_scaled = e / scale_e
    de_scaled = de / scale_de

    mu_i1 = {{k: float(v[0]) for k, v in i1_sets(np.array([e_scaled])).items()}}
    mu_i2 = {{k: float(v[0]) for k, v in i2_sets(np.array([de_scaled])).items()}}

    weights = []
    consequents = []

    for i2_term, i1_term, o_term in rules:
        w = min(mu_i1[i1_term], mu_i2[i2_term])
        weights.append(w)
        consequents.append(C[o_term])

    weights = np.array(weights)
    consequents = np.array(consequents)

    if np.sum(weights) == 0:
        return 0.0

    return np.sum(weights * consequents) / np.sum(weights)


def object_model(t, x, u):
    K, w0, d = 1.0, 1.0, 0.5
    return [x[1], K * u - 2 * d * w0 * x[1] - w0 ** 2 * x[0]]


class PIDController:
    def __init__(self, Kp, Ki, Kd, dt, setpoint=1.0):
        self.Kp, self.Ki, self.Kd = Kp, Ki, Kd
        self.dt, self.setpoint = dt, setpoint
        self.integral, self.prev_error = 0.0, 0.0
        self.first_call = True  # Флаг первого вызова

    def reset(self):
        self.integral, self.prev_error = 0.0, 0.0
        self.first_call = True

    def update(self, measured):
        error = self.setpoint - measured

        if self.first_call:
            derivative = 0.0
            self.first_call = False
        else:
            derivative = (error - self.prev_error) / self.dt

        self.integral += error * self.dt
        output = self.Kp * error + self.Ki * self.integral + self.Kd * derivative
        self.prev_error = error
        return output


class FuzzyMamdaniPID:
    def __init__(self, Ki, dt, setpoint=1.0, scale_e=2.0, scale_de=4.0):
        self.Ki = Ki
        self.dt = dt
        self.setpoint = setpoint
        self.scale_e = scale_e
        self.scale_de = scale_de
        self.integral = 0.0

    def reset(self):
        self.integral = 0.0

    def update(self, measured, de):
        error = self.setpoint - measured
        self.integral += error * self.dt
        u_fuzzy = mamdani_control(error, de, self.scale_e, self.scale_de)
        return u_fuzzy + self.Ki * self.integral


def plot_membership_functions():
    x1 = np.linspace(I1_min, I1_max, 1000)
    x2 = np.linspace(I2_min, I2_max, 1000)
    xo = np.linspace(O_min, O_max, 1000)

    i1 = i1_sets(x1)
    i2 = i2_sets(x2)
    o = o_sets(xo)

    plt.figure(figsize=(10, 4))
    plt.plot(x1, i1["I1NH"], label="I1NH")
    plt.plot(x1, i1["I1NL"], label="I1NL")
    plt.plot(x1, i1["I1Z"], label="I1Z")
    plt.plot(x1, i1["I1PL"], label="I1PL")
    plt.plot(x1, i1["I1PH"], label="I1PH")
    plt.title("Функции принадлежности I1")
    plt.xlabel("I1")
    plt.ylabel("μ")
    plt.grid(True)
    plt.legend()
    plt.tight_layout()

    plt.figure(figsize=(10, 4))
    plt.plot(x2, i2["I2N"], label="I2N")
    plt.plot(x2, i2["I2Z"], label="I2Z")
    plt.plot(x2, i2["I2P"], label="I2P")
    plt.title("Функции принадлежности I2")
    plt.xlabel("I2")
    plt.ylabel("μ")
    plt.grid(True)
    plt.legend()
    plt.tight_layout()

    plt.figure(figsize=(10, 4))
    plt.plot(xo, o["ONH"], label="ONH")
    plt.plot(xo, o["ONL"], label="ONL")
    plt.plot(xo, o["OZ"], label="OZ")
    plt.plot(xo, o["OPL"], label="OPL")
    plt.plot(xo, o["OPH"], label="OPH")
    plt.title("Функции принадлежности O")
    plt.xlabel("O")
    plt.ylabel("μ")
    plt.grid(True)
    plt.legend()
    plt.tight_layout()


def simulate(controller, t_span, dt, x0_state):
    t_eval = np.arange(t_span[0], t_span[1] + dt, dt)
    n = len(t_eval)
    y, u, err = np.zeros(n), np.zeros(n), np.zeros(n)
    state = x0_state
    y[0] = state[0]
    err[0] = controller.setpoint - y[0]

    for i in range(1, n):
        de = -state[1]  # производная ошибки
        u[i] = controller.update(y[i - 1], de) if hasattr(controller, 'scale_e') else controller.update(y[i - 1])

        sol = solve_ivp(lambda t, s: object_model(t, s, u[i]),
                        (t_eval[i - 1], t_eval[i]), state, method='RK45', t_eval=[t_eval[i]])
        state = sol.y[:, -1]
        y[i] = state[0]
        err[i] = controller.setpoint - y[i]

    return t_eval, y, err, u


if __name__ == "__main__":

    print(f"x0 = {{x0}}, y0 = {{y0}}")

    mu_i1 = {{k: float(v[0]) for k, v in i1_sets(np.array([x0])).items()}}
    mu_i2 = {{k: float(v[0]) for k, v in i2_sets(np.array([y0])).items()}}

    print("\\nСтепени принадлежности I1:")
    for k, v in mu_i1.items(): print(f"  {{k}}: {{v:.4f}}")
    print("\\nСтепени принадлежности I2:")
    for k, v in mu_i2.items(): print(f"  {{k}}: {{v:.4f}}")

    z_mamdani = mamdani_control(x0, y0, scale_e=1.0, scale_de=1.0, resolution=1000)
    print(f"\\nРезультат вывода по Мамдани: z0 = {{z_mamdani:.4f}}")

    z_sugeno = sugeno_zero_order_control(x0, y0, scale_e=1.0, scale_de=1.0)
    print(f"\\nРезультат вывода по Сугено: z0 = {{z_sugeno:.4f}}")

    dt = 0.01
    t_span = (0, 10)
    x0_state = [0.0, 0.0]
    setpoint = 1.0

    # 1. Классический ПИД
    pid = PIDController(Kp=3.0, Ki=1.0, Kd=1.5, dt=dt, setpoint=setpoint)
    t_pid, y_pid, err_pid, u_pid = simulate(pid, t_span, dt, x0_state)

    # 2. Нечёткий ПИД (Мамдани)
    fuzzy = FuzzyMamdaniPID(Ki=0.5, dt=dt, setpoint=setpoint, scale_e=2.0, scale_de=4.0)
    t_fuzzy, y_fuzzy, err_fuzzy, u_fuzzy = simulate(fuzzy, t_span, dt, x0_state)

    plot_membership_functions()

    # Графики
    plt.figure(figsize=(12, 8))

    plt.subplot(3, 1, 1)
    plt.plot(t_pid, y_pid, 'b-', label='ПИД')
    plt.plot(t_fuzzy, y_fuzzy, 'r-', label='Нечёткий ПИД (Мамдани)')
    plt.axhline(y=setpoint, linestyle='--', color='gray', label='Задание')
    plt.ylabel('y(t)')
    plt.legend()
    plt.grid(True)

    plt.subplot(3, 1, 2)
    plt.plot(t_pid, err_pid, 'b-', label='Ошибка (ПИД)')
    plt.plot(t_fuzzy, err_fuzzy, 'r-', label='Ошибка (нечёткий)')
    plt.ylabel('Ошибка')
    plt.legend()
    plt.grid(True)

    plt.subplot(3, 1, 3)
    plt.plot(t_pid, u_pid, 'b-', label='Управление (ПИД)')
    plt.plot(t_fuzzy, u_fuzzy, 'r-', label='Управление (нечёткий)')
    plt.xlabel('Время (с)')
    plt.ylabel('u(t)')
    plt.legend()
    plt.grid(True)

    plt.tight_layout()


    def settling_time(t, y, setpoint, tol=0.05):
        band = tol * abs(setpoint)
        for i in range(len(t)):
            if np.all(np.abs(y[i:] - setpoint) <= band):
                return t[i]
        return np.nan


    t_settle_pid = settling_time(t_pid, y_pid, setpoint)
    t_settle_fuzzy = settling_time(t_fuzzy, y_fuzzy, setpoint)
    overshoot_pid = max(0, (np.max(y_pid) - setpoint) / setpoint * 100)
    overshoot_fuzzy = max(0, (np.max(y_fuzzy) - setpoint) / setpoint * 100)

    print(f"ПИД:           время регулирования = {{t_settle_pid:.2f}} с, перерегулирование = {{overshoot_pid:.1f}}%")
    print(f"Нечёткий ПИД:  время регулирования = {{t_settle_fuzzy:.2f}} с, перерегулирование = {{overshoot_fuzzy:.1f}}%")

    plt.show()
'''

    return code


# =========================
# Нечёткий вывод
# =========================
def mamdani_control(e, de, i1_params, i2_params, o_params, resolution=200):
    e_scaled = e
    de_scaled = de

    mu_i1 = {k: float(v[0]) for k, v in i1_sets(np.array([e_scaled]), i1_params).items()}
    mu_i2 = {k: float(v[0]) for k, v in i2_sets(np.array([de_scaled]), i2_params).items()}

    O_min, O_1, O_2, O_3, O_4, O_max = o_params
    z = np.linspace(O_min, O_max, resolution)
    out_mfs = o_sets(z, o_params)

    aggregated = np.zeros_like(z)

    for i2_term, i1_term, o_term in rules:
        alpha = min(mu_i1[i1_term], mu_i2[i2_term])
        if alpha > 0:
            clipped = np.minimum(alpha, out_mfs[o_term])
            aggregated = np.maximum(aggregated, clipped)

    if np.sum(aggregated) == 0:
        return 0.0
    return np.sum(z * aggregated) / np.sum(aggregated)


def sugeno_zero_order_control(e, de, i1_params, i2_params, C):
    e_scaled = e
    de_scaled = de

    mu_i1 = {k: float(v[0]) for k, v in i1_sets(np.array([e_scaled]), i1_params).items()}
    mu_i2 = {k: float(v[0]) for k, v in i2_sets(np.array([de_scaled]), i2_params).items()}

    weights = []
    consequents = []

    for i2_term, i1_term, o_term in rules:
        w = min(mu_i1[i1_term], mu_i2[i2_term])
        weights.append(w)
        consequents.append(C[o_term])

    weights = np.array(weights)
    consequents = np.array(consequents)

    if np.sum(weights) == 0:
        return 0.0
    return np.sum(weights * consequents) / np.sum(weights)


# =========================
# Модель объекта
# =========================
def object_model(t, x, u):
    K, w0, d = 1.0, 1.0, 0.5
    return [x[1], K * u - 2 * d * w0 * x[1] - w0 ** 2 * x[0]]


# =========================
# ПИД регулятор
# =========================
class PIDController:
    def __init__(self, Kp, Ki, Kd, dt, setpoint=1.0):
        self.Kp, self.Ki, self.Kd = Kp, Ki, Kd
        self.dt, self.setpoint = dt, setpoint
        self.integral, self.prev_error = 0.0, 0.0
        self.first_call = True

    def reset(self):
        self.integral, self.prev_error = 0.0, 0.0
        self.first_call = True

    def update(self, measured):
        error = self.setpoint - measured
        if self.first_call:
            derivative = 0.0
            self.first_call = False
        else:
            derivative = (error - self.prev_error) / self.dt
        self.integral += error * self.dt
        output = self.Kp * error + self.Ki * self.integral + self.Kd * derivative
        self.prev_error = error
        return output


# =========================
# Нечёткий ПИД (Мамдани)
# =========================
class FuzzyMamdaniPID:
    def __init__(self, Ki, dt, setpoint=1.0, i1_params=None, i2_params=None, o_params=None):
        self.Ki = Ki
        self.dt = dt
        self.setpoint = setpoint
        self.i1_params = i1_params
        self.i2_params = i2_params
        self.o_params = o_params
        self.integral = 0.0

    def reset(self):
        self.integral = 0.0

    def update(self, measured, de):
        error = self.setpoint - measured
        self.integral += error * self.dt
        u_fuzzy = mamdani_control(error, de, self.i1_params, self.i2_params, self.o_params)
        return u_fuzzy + self.Ki * self.integral


# =========================
# Симуляция
# =========================
def simulate(controller, t_span, dt, x0_state, is_fuzzy=False):
    t_eval = np.arange(t_span[0], t_span[1] + dt, dt)
    n = len(t_eval)
    y, u, err = np.zeros(n), np.zeros(n), np.zeros(n)
    state = x0_state.copy()
    y[0] = state[0]
    err[0] = controller.setpoint - y[0]

    for i in range(1, n):
        if is_fuzzy:
            de = -state[1]
            u[i] = controller.update(y[i - 1], de)
        else:
            u[i] = controller.update(y[i - 1])

        sol = solve_ivp(lambda t, s: object_model(t, s, u[i]),
                        (t_eval[i - 1], t_eval[i]), state, method='RK45', t_eval=[t_eval[i]])
        state = sol.y[:, -1]
        y[i] = state[0]
        err[i] = controller.setpoint - y[i]

    return t_eval, y, err, u


def settling_time(t, y, setpoint, tol=0.05):
    band = tol * abs(setpoint)
    for i in range(len(t)):
        if np.all(np.abs(y[i:] - setpoint) <= band):
            return t[i]
    return np.nan


# =========================
# Функции для построения графиков
# =========================
def plot_membership_functions(i1_params, i2_params, o_params):
    I1_min, I1_1, I1_2, I1_3, I1_4, I1_max = i1_params
    I2_min, I2_1, I2_2, I2_max = i2_params
    O_min, O_1, O_2, O_3, O_4, O_max = o_params

    x1 = np.linspace(I1_min, I1_max, 1000)
    x2 = np.linspace(I2_min, I2_max, 1000)
    xo = np.linspace(O_min, O_max, 1000)

    i1 = i1_sets(x1, i1_params)
    i2 = i2_sets(x2, i2_params)
    o = o_sets(xo, o_params)

    fig1, ax1 = plt.subplots(figsize=(10, 4))
    ax1.plot(x1, i1["I1NH"], label="I1NH")
    ax1.plot(x1, i1["I1NL"], label="I1NL")
    ax1.plot(x1, i1["I1Z"], label="I1Z")
    ax1.plot(x1, i1["I1PL"], label="I1PL")
    ax1.plot(x1, i1["I1PH"], label="I1PH")
    ax1.set_title("Функции принадлежности I1")
    ax1.set_xlabel("I1")
    ax1.set_ylabel("μ")
    ax1.grid(True)
    ax1.legend()

    fig2, ax2 = plt.subplots(figsize=(10, 4))
    ax2.plot(x2, i2["I2N"], label="I2N")
    ax2.plot(x2, i2["I2Z"], label="I2Z")
    ax2.plot(x2, i2["I2P"], label="I2P")
    ax2.set_title("Функции принадлежности I2")
    ax2.set_xlabel("I2")
    ax2.set_ylabel("μ")
    ax2.grid(True)
    ax2.legend()

    fig3, ax3 = plt.subplots(figsize=(10, 4))
    ax3.plot(xo, o["ONH"], label="ONH")
    ax3.plot(xo, o["ONL"], label="ONL")
    ax3.plot(xo, o["OZ"], label="OZ")
    ax3.plot(xo, o["OPL"], label="OPL")
    ax3.plot(xo, o["OPH"], label="OPH")
    ax3.set_title("Функции принадлежности O")
    ax3.set_xlabel("O")
    ax3.set_ylabel("μ")
    ax3.grid(True)
    ax3.legend()

    return fig1, fig2, fig3


def plot_simulation_results(t_pid, y_pid, err_pid, u_pid, t_fuzzy, y_fuzzy, err_fuzzy, u_fuzzy, setpoint):
    fig, axes = plt.subplots(3, 1, figsize=(12, 10))

    axes[0].plot(t_pid, y_pid, 'b-', label='ПИД')
    axes[0].plot(t_fuzzy, y_fuzzy, 'r-', label='Нечёткий ПИД (Мамдани)')
    axes[0].axhline(y=setpoint, linestyle='--', color='gray', label='Задание')
    axes[0].set_ylabel('y(t)')
    axes[0].legend()
    axes[0].grid(True)

    axes[1].plot(t_pid, err_pid, 'b-', label='Ошибка (ПИД)')
    axes[1].plot(t_fuzzy, err_fuzzy, 'r-', label='Ошибка (нечёткий)')
    axes[1].set_ylabel('Ошибка')
    axes[1].legend()
    axes[1].grid(True)

    axes[2].plot(t_pid, u_pid, 'b-', label='Управление (ПИД)')
    axes[2].plot(t_fuzzy, u_fuzzy, 'r-', label='Управление (нечёткий)')
    axes[2].set_xlabel('Время (с)')
    axes[2].set_ylabel('u(t)')
    axes[2].legend()
    axes[2].grid(True)

    plt.tight_layout()
    return fig


# =========================
# Генерация отчёта Word
# =========================

TEMPLATE_PATH = "pattern_ai_laba2.docx"


def fig_to_bytes(fig):
    """Переводит matplotlib Figure в BytesIO для вставки в Word."""
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img


def replace_text_in_paragraph(paragraph, replacements):
    """
    Заменяет {{метки}} в абзаце.
    Работает даже если Word разбил метку на несколько run-ов.
    """
    full_text = "".join(run.text for run in paragraph.runs)

    changed = False
    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, str(value))
            changed = True

    if changed:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)


def replace_text_everywhere(doc, replacements):
    """Заменяет текстовые метки во всех абзацах и таблицах документа."""
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def replace_placeholder_with_image(doc, placeholder, image_stream, width_inches=5.8):
    """Заменяет абзац с меткой на картинку."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(width_inches))
            return

    # Если метка вдруг оказалась внутри таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(width_inches))
                        return


def replace_placeholder_with_code(doc, placeholder, code_string):
    """Заменяет метку на текст программного кода с сохранением форматирования."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Очищаем абзац
            paragraph.clear()
            # Добавляем текст кода
            run = paragraph.add_run(code_string)
            # Можно добавить настройки шрифта для кода (например, Courier New)
            run.font.name = 'Courier New'
            run.font.size = Inches(0.1)  # ~10pt
            return

    # Если метка в таблице
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.clear()
                        run = paragraph.add_run(code_string)
                        run.font.name = 'Courier New'
                        run.font.size = Inches(0.1)
                        return


def generate_report_docx(
        fio,
        variant,
        data,
        C,
        mu_i1,
        mu_i2,
        z_mamdani,
        z_sugeno,
        fig1,
        fig2,
        fig3,
        fig_sim,
        t_settle_pid,
        t_settle_fuzzy,
        overshoot_pid,
        overshoot_fuzzy,
        code_string
):
    """Создаёт готовый отчёт Word на основе шаблона."""

    doc = Document(TEMPLATE_PATH)

    i1_params = data["I1"]
    i2_params = data["I2"]
    o_params = data["O"]
    x0 = data["x0"]
    y0 = data["y0"]

    replacements = {
        "{{name}}": fio,
        "{{var}}": variant,

        "{{I1_min}}": i1_params[0],
        "{{I1_1}}": i1_params[1],
        "{{I1_2}}": i1_params[2],
        "{{I1_3}}": i1_params[3],
        "{{I1_4}}": i1_params[4],
        "{{I1_max}}": i1_params[5],

        "{{I2_min}}": i2_params[0],
        "{{I2_1}}": i2_params[1],
        "{{I2_2}}": i2_params[2],
        "{{I2_max}}": i2_params[3],

        "{{O_min}}": o_params[0],
        "{{O_1}}": o_params[1],
        "{{O_2}}": o_params[2],
        "{{O_3}}": o_params[3],
        "{{O_4}}": o_params[4],
        "{{O_max}}": o_params[5],

        "{{O_NH}}": C["ONH"],
        "{{O_NL}}": C["ONL"],
        "{{O_Z}}": C["OZ"],
        "{{O_PL}}": C["OPL"],
        "{{O_PH}}": C["OPH"],

        "{{x0}}": x0,
        "{{y0}}": y0,

        "{{I1NH}}": f"{mu_i1['I1NH']:.4f}",
        "{{I1NL}}": f"{mu_i1['I1NL']:.4f}",
        "{{I1Z}}": f"{mu_i1['I1Z']:.4f}",
        "{{I1PL}}": f"{mu_i1['I1PL']:.4f}",
        "{{I1PH}}": f"{mu_i1['I1PH']:.4f}",

        "{{I2N}}": f"{mu_i2['I2N']:.4f}",
        "{{I2Z}}": f"{mu_i2['I2Z']:.4f}",
        "{{I2P}}": f"{mu_i2['I2P']:.4f}",

        "{{z0_mam}}": f"{z_mamdani:.4f}",
        "{{z0_sug}}": f"{z_sugeno:.4f}",

        "{{time_pid}}": f"{t_settle_pid:.2f} с",
        "{{time_n}}": f"{t_settle_fuzzy:.2f} с",
        "{{pere_pid}}": f"{overshoot_pid:.1f} %",
        "{{pere_n}}": f"{overshoot_fuzzy:.1f} %",
    }

    replace_text_everywhere(doc, replacements)

    replace_placeholder_with_image(doc, "{{func_i1}}", fig_to_bytes(fig1), width_inches=5.8)
    replace_placeholder_with_image(doc, "{{func_i2}}", fig_to_bytes(fig2), width_inches=5.8)
    replace_placeholder_with_image(doc, "{{func_o}}", fig_to_bytes(fig3), width_inches=5.8)
    replace_placeholder_with_image(doc, "{{func_pid}}", fig_to_bytes(fig_sim), width_inches=6.2)
    replace_placeholder_with_code(doc, "{{code}}", code_string)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


# =========================
# Streamlit приложение
# =========================
def run():
    st.set_page_config(page_title="Нечёткие регуляторы", layout="wide")
    st.title("Автоматизированная система расчёта нечётких регуляторов")
    st.markdown("---")

    # Выбор варианта
    col1, col2 = st.columns([1, 2])
    with col1:
        variant = st.selectbox("Выберите вариант", list(range(1, 36)), index=12)  # index=12 для варианта 13
        fio = st.text_input("Введите ФИО студента", value="")
        st.write(f"**Выбран вариант {variant}**")

    # Получение данных варианта
    data = variants_data[variant]
    i1_params = data["I1"]
    i2_params = data["I2"]
    o_params = data["O"]
    C = {
        "ONH": data["C"][0],
        "ONL": data["C"][1],
        "OZ": data["C"][2],
        "OPL": data["C"][3],
        "OPH": data["C"][4]
    }
    x0, y0 = data["x0"], data["y0"]

    with col2:
        st.write("### Параметры текущего варианта")
        col2a, col2b, col2c = st.columns(3)
        with col2a:
            st.write("**I1 параметры:**")
            st.write(
                f"min={i1_params[0]}, 1={i1_params[1]}, 2={i1_params[2]}, 3={i1_params[3]}, 4={i1_params[4]}, max={i1_params[5]}")
        with col2b:
            st.write("**I2 параметры:**")
            st.write(f"min={i2_params[0]}, 1={i2_params[1]}, 2={i2_params[2]}, max={i2_params[3]}")
        with col2c:
            st.write("**O параметры:**")
            st.write(
                f"min={o_params[0]}, 1={o_params[1]}, 2={o_params[2]}, 3={o_params[3]}, 4={o_params[4]}, max={o_params[5]}")
        st.write(f"**Входные значения:** x0 = {x0}, y0 = {y0}")
        st.write(f"**Коэффициенты Sugeno:** {C}")

    st.markdown("---")

    # Расчёт для заданных x0, y0
    st.header("1. Расчёт для заданных входных значений")

    mu_i1 = {k: float(v[0]) for k, v in i1_sets(np.array([x0]), i1_params).items()}
    mu_i2 = {k: float(v[0]) for k, v in i2_sets(np.array([y0]), i2_params).items()}

    col1, col2 = st.columns(2)
    with col1:
        st.write("**Степени принадлежности I1:**")
        for k, v in mu_i1.items():
            st.write(f"  {k}: {v:.4f}")
    with col2:
        st.write("**Степени принадлежности I2:**")
        for k, v in mu_i2.items():
            st.write(f"  {k}: {v:.4f}")

    # Расчёт по Мамдани и Сугено
    z_mamdani = mamdani_control(x0, y0, i1_params, i2_params, o_params, resolution=1000)
    z_sugeno = sugeno_zero_order_control(x0, y0, i1_params, i2_params, C)

    st.write("### Результаты нечёткого вывода")
    col1, col2 = st.columns(2)
    with col1:
        st.success(f"**Метод Мамдани:** z0 = {z_mamdani:.4f}")
    with col2:
        st.success(f"**Метод Сугено (0-го порядка):** z0 = {z_sugeno:.4f}")

    st.markdown("---")

    # Функции принадлежности
    st.header("2. Функции принадлежности")
    fig1, fig2, fig3 = plot_membership_functions(i1_params, i2_params, o_params)
    col1, col2 = st.columns(2)
    with col1:
        st.pyplot(fig1)
        st.pyplot(fig2)
    with col2:
        st.pyplot(fig3)

    st.markdown("---")

    # Симуляция системы управления
    st.header("3. Моделирование системы управления")

    dt = 0.01
    t_span = (0, 10)
    x0_state = [0.0, 0.0]
    setpoint = 1.0

    # Параметры ПИД регулятора (можно настроить)
    Kp, Ki, Kd = 3.0, 1.0, 1.5

    with st.expander("Настройка параметров регуляторов"):
        col1, col2 = st.columns(2)
        with col1:
            st.write("**ПИД регулятор**")
            Kp = st.number_input("Kp", value=3.0, step=0.5)
            Ki = st.number_input("Ki", value=1.0, step=0.5)
            Kd = st.number_input("Kd", value=1.5, step=0.5)
        with col2:
            st.write("**Нечёткий ПИД**")
            fuzzy_Ki = st.number_input("Ki (нечёткий)", value=0.5, step=0.1)

    # Запуск симуляции
    if st.button("Выполнить моделирование", type="primary"):
        with st.spinner("Моделирование..."):
            # ПИД регулятор
            pid = PIDController(Kp=Kp, Ki=Ki, Kd=Kd, dt=dt, setpoint=setpoint)
            t_pid, y_pid, err_pid, u_pid = simulate(pid, t_span, dt, x0_state, is_fuzzy=False)

            # Нечёткий ПИД
            fuzzy = FuzzyMamdaniPID(Ki=fuzzy_Ki, dt=dt, setpoint=setpoint,
                                    i1_params=i1_params, i2_params=i2_params, o_params=o_params)
            t_fuzzy, y_fuzzy, err_fuzzy, u_fuzzy = simulate(fuzzy, t_span, dt, x0_state, is_fuzzy=True)

            # Расчёт метрик
            t_settle_pid = settling_time(t_pid, y_pid, setpoint)
            t_settle_fuzzy = settling_time(t_fuzzy, y_fuzzy, setpoint)
            overshoot_pid = max(0, (np.max(y_pid) - setpoint) / setpoint * 100)
            overshoot_fuzzy = max(0, (np.max(y_fuzzy) - setpoint) / setpoint * 100)

            # Вывод метрик
            st.write("### Сравнительные метрики")
            col1, col2 = st.columns(2)
            with col1:
                st.info(f"**ПИД:** время регулирования = {t_settle_pid:.2f} с, перерегулирование = {overshoot_pid:.1f}%")
            with col2:
                st.info(
                    f"**Нечёткий ПИД:** время регулирования = {t_settle_fuzzy:.2f} с, перерегулирование = {overshoot_fuzzy:.1f}%")

            # Графики
            fig_sim = plot_simulation_results(t_pid, y_pid, err_pid, u_pid,
                                              t_fuzzy, y_fuzzy, err_fuzzy, u_fuzzy, setpoint)

            # Генерация кода для варианта
            code_string = generate_code_for_variant(variant, data, C)

            if fio.strip():
                report_file = generate_report_docx(
                    fio=fio,
                    variant=variant,
                    data=data,
                    C=C,
                    mu_i1=mu_i1,
                    mu_i2=mu_i2,
                    z_mamdani=z_mamdani,
                    z_sugeno=z_sugeno,
                    fig1=fig1,
                    fig2=fig2,
                    fig3=fig3,
                    fig_sim=fig_sim,
                    t_settle_pid=t_settle_pid,
                    t_settle_fuzzy=t_settle_fuzzy,
                    overshoot_pid=overshoot_pid,
                    overshoot_fuzzy=overshoot_fuzzy,
                    code_string=code_string
                )

                st.download_button(
                    label="Скачать готовый отчёт Word",
                    data=report_file,
                    file_name=f"report_laba2_variant_{variant}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Введите ФИО, чтобы сформировать отчёт.")
            st.pyplot(fig_sim)

            # Вывод результатов в таблицу
            st.write("### Результаты в табличной форме")
            results_df = pd.DataFrame({
                'Время (с)': t_pid[::100],
                'PID y(t)': y_pid[::100],
                'Fuzzy y(t)': y_fuzzy[::100],
                'PID ошибка': err_pid[::100],
                'Fuzzy ошибка': err_fuzzy[::100]
            })
            st.dataframe(results_df)

    st.markdown("---")
    st.caption("Разработано для автоматизации расчётов нечётких регуляторов по индивидуальному заданию №2")
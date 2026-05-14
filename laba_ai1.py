# app.py
import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import random
import math
from dataclasses import dataclass
from typing import List, Tuple, Dict
import io
from docx import Document
from docx.shared import Inches, Pt
import zipfile
import tempfile
import os
import re
from pathlib import Path
st.set_page_config(layout="wide", page_title="Генетические алгоритмы", page_icon="🧬")


# Инициализация session_state для хранения результатов
if 'task1_results' not in st.session_state:
    st.session_state.task1_results = None
if 'task1_ga' not in st.session_state:
    st.session_state.task1_ga = None
if 'task1_populations' not in st.session_state:
    st.session_state.task1_populations = None
if 'task1_avg_fitness' not in st.session_state:
    st.session_state.task1_avg_fitness = None
if 'task1_figures' not in st.session_state:
    st.session_state.task1_figures = {}
if 'task2_results' not in st.session_state:
    st.session_state.task2_results = None
if 'task2_figures' not in st.session_state:
    st.session_state.task2_figures = {}
# ============================================================
# БЛОК ЗАДАНИЯ №1 (базовое) - Поиск экстремума функции
# ============================================================

class Task1GA:
    def __init__(self, variant: int):
        self.variant = variant
        self.set_params_by_variant()
        self.MAX_INT = 2 ** self.chromosome_length - 1

    def set_params_by_variant(self):
        """Установка параметров по варианту (таблица 1, область [-5;5])"""
        variants_5 = {
            1: {"func": lambda x: 2 * x - 3.3 * x * math.cos(9.3 * x) - 2.3 * math.sin(1.7 * x),
                "extreme": "min", "coding": "binary", "pc": 0.9, "pm": 0.2,
                "parent_select": "tournament", "reduction": "tournament"},
            2: {"func": lambda x: math.exp(-0.5 * x ** 2) * (1 if math.cos(9.3 * x - 1) > 0 else -1),
                "extreme": "min", "coding": "gray", "pc": 0.9, "pm": 0.1,
                "parent_select": "random", "reduction": "roulette"},
            3: {"func": lambda x: -0.6 * x + 5.3 * abs(math.cos(6.1 * x)) * math.cos(3.6 * x),
                "extreme": "min", "coding": "binary", "pc": 0.8, "pm": 0.2,
                "parent_select": "random", "reduction": "tournament"},
            4: {"func": lambda x: -1.3 * math.sin(1.6 * x ** 2 - 0.3) * math.exp(-0.3 * x + 0.5),
                "extreme": "max", "coding": "gray", "pc": 0.85, "pm": 0.15,
                "parent_select": "best_with_best", "reduction": "roulette"},
            5: {"func": lambda x: 0.9 * abs(math.sin(3.7 * x)) * math.cos(6.1 * x),
                "extreme": "max", "coding": "binary", "pc": 0.9, "pm": 0.15,
                "parent_select": "random", "reduction": "tournament"},
            6: {"func": lambda x: 0.8 * x + 1.4 * math.cos(1.8 * x ** 2) * math.exp(0.4 * x),
                "extreme": "min", "coding": "gray", "pc": 0.85, "pm": 0.1,
                "parent_select": "best_with_best", "reduction": "roulette"},
            7: {"func": lambda x: -math.sin(0.9 * x - 1) - math.sin(1.8 * x - 1) * math.cos(7.8 * x),
                "extreme": "max", "coding": "binary", "pc": 0.8, "pm": 0.15,
                "parent_select": "best_with_best", "reduction": "tournament"},
            8: {"func": lambda x: 0.8 * x + 1.1 * x * math.sin(9.3 * x) - 0.7 * math.cos(0.8 * x),
                "extreme": "min", "coding": "gray", "pc": 0.9, "pm": 0.1,
                "parent_select": "random", "reduction": "roulette"},
            9: {"func": lambda x: 0.2 * x - 1.1 * math.exp(-0.4 * x ** 2) * (1 if math.cos(9.5 * x + 1.5) > 0 else -1),
                "extreme": "min", "coding": "binary", "pc": 0.85, "pm": 0.15,
                "parent_select": "best_with_best", "reduction": "tournament"},
            10: {"func": lambda x: 0.3 * x + x * math.cos(7.3 * x) - 0.7 * math.sin(1.3 * x),
                 "extreme": "max", "coding": "gray", "pc": 0.8, "pm": 0.2,
                 "parent_select": "random", "reduction": "roulette"},
        }

        # Варианты для области [-1;1] (11-16)
        variants_1 = {
            11: {"func": lambda x: 5.1 * abs(math.cos(15 * x)) * math.cos(33 * x),
                 "extreme": "max", "coding": "binary", "pc": 0.85, "pm": 0.15,
                 "parent_select": "random", "reduction": "roulette", "x_min": -1, "x_max": 1},
            12: {"func": lambda x: 3.5 * x + math.sin(47 * x ** 2 + 2) - 6 * (x ** 2),
                 "extreme": "max", "coding": "gray", "pc": 0.8, "pm": 0.1,
                 "parent_select": "best_with_best", "reduction": "tournament", "x_min": -1, "x_max": 1},
            13: {"func": lambda x: 9 * x - 9.7 * abs(math.sin(37 * x)) * math.cos(27 * x),
                 "extreme": "min", "coding": "binary", "pc": 0.9, "pm": 0.2,
                 "parent_select": "random", "reduction": "roulette", "x_min": -1, "x_max": 1},
            14: {"func": lambda x: 3.1 * x - 3 * math.sin(31.4 * x ** 2 + 7) - 6 * (x ** 2),
                 "extreme": "max", "coding": "gray", "pc": 0.85, "pm": 0.2,
                 "parent_select": "random", "reduction": "tournament", "x_min": -1, "x_max": 1},
            15: {"func": lambda x: 6.9 * x + 7.5 * abs(math.cos(39 * x)) * math.cos(31 * x),
                 "extreme": "max", "coding": "binary", "pc": 0.8, "pm": 0.1,
                 "parent_select": "random", "reduction": "roulette", "x_min": -1, "x_max": 1},
            16: {"func": lambda x: 0.7 * x + 1.4 * x * math.sin(43 * x) - 2.3 * math.cos(5.4 * x),
                 "extreme": "max", "coding": "gray", "pc": 0.9, "pm": 0.15,
                 "parent_select": "best_with_best", "reduction": "tournament", "x_min": -1, "x_max": 1},
        }

        # Варианты для области [-3;3] (17-23)
        variants_3 = {
            17: {"func": lambda x: math.cos(1.2 * x - 2) - math.cos(1.7 * x - 1) * math.sin(8.4 * x),
                 "extreme": "max", "coding": "binary", "pc": 0.85, "pm": 0.15,
                 "parent_select": "random", "reduction": "roulette", "x_min": -3, "x_max": 3},
            18: {"func": lambda x: -0.35 * x + 1.2 * math.sin(2.5 * x ** 2) * math.exp(0.3 * x),
                 "extreme": "max", "coding": "gray", "pc": 0.9, "pm": 0.1,
                 "parent_select": "best_with_best", "reduction": "tournament", "x_min": -3, "x_max": 3},
            19: {"func": lambda x: -0.5 * x - 0.7 * x * math.cos(9.2 * x + 3) + 0.9 * math.sin(0.9 * x - 1),
                 "extreme": "min", "coding": "binary", "pc": 0.8, "pm": 0.2,
                 "parent_select": "random", "reduction": "roulette", "x_min": -3, "x_max": 3},
            20: {"func": lambda x: -0.9 * abs(math.cos(4.1 * x)) * math.sin(6.8 * x),
                 "extreme": "min", "coding": "gray", "pc": 0.85, "pm": 0.15,
                 "parent_select": "random", "reduction": "tournament", "x_min": -3, "x_max": 3},
            21: {
                "func": lambda x: -0.5 * x + 1.2 * math.exp(-0.5 * x ** 2) * (1 if math.sin(9.7 * x + 1.3) > 0 else -1),
                "extreme": "max", "coding": "binary", "pc": 0.8, "pm": 0.1,
                "parent_select": "best_with_best", "reduction": "roulette", "x_min": -3, "x_max": 3},
            22: {"func": lambda x: -0.9 * x + 1.5 * math.cos(2.3 * x ** 2 - 0.5) * math.exp(-0.4 * x + 0.3),
                 "extreme": "min", "coding": "gray", "pc": 0.9, "pm": 0.15,
                 "parent_select": "random", "reduction": "tournament", "x_min": -3, "x_max": 3},
            23: {"func": lambda x: 0.2 * x + x * math.sin(8.9 * x) + 0.9 * math.cos(1.5 * x - 2),
                 "extreme": "min", "coding": "binary", "pc": 0.85, "pm": 0.15,
                 "parent_select": "random", "reduction": "roulette", "x_min": -3, "x_max": 3},
        }

        # Определяем область поиска по варианту
        if 1 <= self.variant <= 10:
            self.x_min, self.x_max = -5, 5
            params = variants_5.get(self.variant)
        elif 11 <= self.variant <= 16:
            self.x_min, self.x_max = -1, 1
            params = variants_1.get(self.variant)
        elif 17 <= self.variant <= 23:
            self.x_min, self.x_max = -3, 3
            params = variants_3.get(self.variant)
        else:
            raise ValueError(f"Неверный вариант: {self.variant}")

        if params is None:
            raise ValueError(f"Вариант {self.variant} не найден")

        self.target_func = params["func"]
        self.extreme = params["extreme"]
        self.coding = params["coding"]
        self.pc = params["pc"]
        self.pm = params["pm"]
        self.parent_select = params["parent_select"]
        self.reduction = params["reduction"]

        # Общие параметры
        self.population_size = 50
        self.chromosome_length = 20
        self.generations = 30
        self.tournament_size = 3

    def int_to_binary(self, value: int) -> str:
        return format(value, f'0{self.chromosome_length}b')

    def binary_to_int(self, binary: str) -> int:
        return int(binary, 2)

    def gray_to_binary(self, gray: str) -> str:
        binary = gray[0]
        for i in range(1, len(gray)):
            binary += str(int(binary[i - 1]) ^ int(gray[i]))
        return binary

    def binary_to_gray(self, binary: str) -> str:
        gray = binary[0]
        for i in range(1, len(binary)):
            gray += str(int(binary[i - 1]) ^ int(binary[i]))
        return gray

    def code(self, value: int) -> str:
        binary = self.int_to_binary(value)
        if self.coding == "gray":
            return self.binary_to_gray(binary)
        return binary

    def decode(self, chromosome: str) -> int:
        if self.coding == "gray":
            binary = self.gray_to_binary(chromosome)
            return self.binary_to_int(binary)
        return self.binary_to_int(chromosome)

    def int_to_x(self, value: int) -> float:
        max_int = self.MAX_INT
        return self.x_min + value * (self.x_max - self.x_min) / max_int

    def get_fitness_shift(self) -> float:
        xs = np.linspace(self.x_min, self.x_max, 5000)
        ys = [self.target_func(x) for x in xs]
        if self.extreme == "min":
            return max(ys)
        else:
            return abs(min(ys))

    def fitness(self, x: float) -> float:
        f_val = self.target_func(x)
        if self.extreme == "min":
            return self.fitness_shift - f_val + 1e-9
        else:
            return f_val - self.fitness_shift + 1e-9

    @dataclass
    class Individual:
        chromosome: str
        parent_ref: 'Task1GA' = None

        @property
        def x(self) -> float:
            return self.parent_ref.int_to_x(self.parent_ref.decode(self.chromosome))

        @property
        def f(self) -> float:
            return self.parent_ref.target_func(self.x)

        @property
        def fitness(self) -> float:
            return self.parent_ref.fitness(self.x)

    def generate_random_individual(self):
        value = random.randint(0, self.MAX_INT)
        chromosome = self.code(value)
        ind = self.Individual(chromosome)
        ind.parent_ref = self
        return ind

    def generate_initial_population(self):
        return [self.generate_random_individual() for _ in range(self.population_size)]

    def form_parent_pairs(self, population):
        if self.parent_select == "random":
            shuffled = population[:]
            random.shuffle(shuffled)
            pairs = []
            for i in range(0, len(shuffled) - 1, 2):
                pairs.append((shuffled[i], shuffled[i + 1]))
            return pairs
        elif self.parent_select == "best_with_best":
            sorted_pop = sorted(population, key=lambda ind: ind.fitness, reverse=True)
            pairs = []
            for i in range(0, len(sorted_pop) - 1, 2):
                pairs.append((sorted_pop[i], sorted_pop[i + 1]))
            random.shuffle(pairs)
            return pairs
        elif self.parent_select == "tournament":
            pairs = []
            for _ in range(self.population_size // 2):
                p1 = self.tournament_select(population)
                p2 = self.tournament_select(population)
                pairs.append((p1, p2))
            return pairs
        return []

    def tournament_select(self, population):
        tournament = random.sample(population, self.tournament_size)
        return max(tournament, key=lambda ind: ind.fitness)

    def crossover(self, p1, p2):
        if random.random() > self.pc:
            return self.Individual(p1.chromosome, self), self.Individual(p2.chromosome, self)
        point = random.randint(1, self.chromosome_length - 1)
        c1 = p1.chromosome[:point] + p2.chromosome[point:]
        c2 = p2.chromosome[:point] + p1.chromosome[point:]
        return self.Individual(c1, self), self.Individual(c2, self)

    def mutate(self, individual):
        if random.random() > self.pm:
            return individual
        chrom = list(individual.chromosome)
        bit = random.randint(0, self.chromosome_length - 1)
        chrom[bit] = '1' if chrom[bit] == '0' else '0'
        return self.Individual(''.join(chrom), self)

    def reduction_tournament(self, population, offspring):
        candidates = population + offspring
        new_pop = []
        for _ in range(self.population_size):
            new_pop.append(self.tournament_select(candidates))
        return new_pop

    def reduction_roulette(self, population, offspring):
        candidates = population + offspring
        fitness_sum = sum(ind.fitness for ind in candidates)
        if fitness_sum <= 0:
            return random.sample(candidates, self.population_size)
        probs = [ind.fitness / fitness_sum for ind in candidates]
        indices = np.random.choice(len(candidates), size=self.population_size, replace=True, p=probs)
        return [candidates[i] for i in indices]

    def run(self, seed=42):
        random.seed(seed)
        np.random.seed(seed)

        self.fitness_shift = self.get_fitness_shift()

        population = self.generate_initial_population()
        populations = {0: population[:]}
        avg_fitness_history = [sum(ind.fitness for ind in population) / self.population_size]

        for gen in range(1, self.generations + 1):
            parent_pairs = self.form_parent_pairs(population)
            offspring = []
            for p1, p2 in parent_pairs:
                c1, c2 = self.crossover(p1, p2)
                c1 = self.mutate(c1)
                c2 = self.mutate(c2)
                offspring.extend([c1, c2])

            if self.reduction == "tournament":
                population = self.reduction_tournament(population, offspring)
            else:
                population = self.reduction_roulette(population, offspring)

            if gen == 3:
                populations[3] = population[:]
            if gen == self.generations:
                populations[self.generations] = population[:]

            avg_fitness_history.append(sum(ind.fitness for ind in population) / self.population_size)

        return populations, avg_fitness_history

def normalize_docx_placeholders(template_bytes: bytes) -> bytes:
    """Склеивает плейсхолдеры вида {{...}}, если Word разорвал их на несколько runs.
    Это важно для вашего шаблона: часть меток, например {{pop_3_x_50}}, разрезана переносами/тегами XML.
    """
    src = io.BytesIO(template_bytes)
    out = io.BytesIO()

    def normalize_marker(match: re.Match) -> str:
        raw = match.group(0)
        raw = re.sub(r"<[^>]+>", "", raw)       # убрать XML-теги внутри метки
        raw = re.sub(r"\s+", "", raw)           # убрать пробелы/переносы внутри метки
        return raw

    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                xml = re.sub(r"\{\{.*?\}\}", normalize_marker, xml, flags=re.S)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    out.seek(0)
    return out.getvalue()


def replace_text_in_doc(doc: Document, values: dict):
    """Заменяет текстовые метки в абзацах и таблицах."""
    def replace_in_paragraph(paragraph):
        for run in paragraph.runs:
            for key, value in values.items():
                marker = "{{" + key + "}}"
                if marker in run.text:
                    run.text = run.text.replace(marker, str(value))

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def get_program_code_for_report(ga) -> str:
    """Генерирует корректный самостоятельный Python-код для отчета под конкретный вариант."""
    from string import Template
    import textwrap

    python_funcs = {
        1: "2*x - 3.3*x*math.cos(9.3*x) - 2.3*math.sin(1.7*x)",
        2: "math.exp(-0.5*x**2) * (1 if math.cos(9.3*x - 1) > 0 else -1)",
        3: "-0.6*x + 5.3*abs(math.cos(6.1*x))*math.cos(3.6*x)",
        4: "-1.3*math.sin(1.6*x**2 - 0.3)*math.exp(-0.3*x + 0.5)",
        5: "0.9*abs(math.sin(3.7*x))*math.cos(6.1*x)",
        6: "0.8*x + 1.4*math.cos(1.8*x**2)*math.exp(0.4*x)",
        7: "-math.sin(0.9*x - 1) - math.sin(1.8*x - 1)*math.cos(7.8*x)",
        8: "0.8*x + 1.1*x*math.sin(9.3*x) - 0.7*math.cos(0.8*x)",
        9: "0.2*x - 1.1*math.exp(-0.4*x**2) * (1 if math.cos(9.5*x + 1.5) > 0 else -1)",
        10: "0.3*x + x*math.cos(7.3*x) - 0.7*math.sin(1.3*x)",
        11: "5.1*abs(math.cos(15*x))*math.cos(33*x)",
        12: "3.5*x + math.sin(47*x**2 + 2) - 6*x**2",
        13: "9*x - 9.7*abs(math.sin(37*x))*math.cos(27*x)",
        14: "3.1*x - 3*math.sin(31.4*x**2 + 7) - 6*x**2",
        15: "6.9*x + 7.5*abs(math.cos(39*x))*math.cos(31*x)",
        16: "0.7*x + 1.4*x*math.sin(43*x) - 2.3*math.cos(5.4*x)",
        17: "math.cos(1.2*x - 2) - math.cos(1.7*x - 1)*math.sin(8.4*x)",
        18: "-0.35*x + 1.2*math.sin(2.5*x**2)*math.exp(0.3*x)",
        19: "-0.5*x - 0.7*x*math.cos(9.2*x + 3) + 0.9*math.sin(0.9*x - 1)",
        20: "-0.9*abs(math.cos(4.1*x))*math.sin(6.8*x)",
        21: "-0.5*x + 1.2*math.exp(-0.5*x**2) * (1 if math.sin(9.7*x + 1.3) > 0 else -1)",
        22: "-0.9*x + 1.5*math.cos(2.3*x**2 - 0.5)*math.exp(-0.4*x + 0.3)",
        23: "0.2*x + x*math.sin(8.9*x) + 0.9*math.cos(1.5*x - 2)",
    }
    func_code = python_funcs.get(ga.variant)
    if not func_code:
        raise ValueError(f"Не найдена функция для варианта {ga.variant}")

    extreme_text = "минимума" if ga.extreme == "min" else "максимума"
    parent_select_text = {
        "random": "случайный отбор",
        "best_with_best": "отбор лучших с лучшими",
        "tournament": "турнирный отбор",
    }.get(ga.parent_select, ga.parent_select)
    reduction_text = {
        "tournament": "турнирная редукция",
        "roulette": "редукция рулеткой",
    }.get(ga.reduction, ga.reduction)
    coding_text = "двоичный код" if ga.coding == "binary" else "код Грея"

    if ga.coding == "gray":
        coding_code = '''
def gray_to_binary(gray: str) -> str:
    binary = gray[0]
    for i in range(1, len(gray)):
        binary += str(int(binary[i - 1]) ^ int(gray[i]))
    return binary


def binary_to_gray(binary: str) -> str:
    gray = binary[0]
    for i in range(1, len(binary)):
        gray += str(int(binary[i - 1]) ^ int(binary[i]))
    return gray


def encode(value: int) -> str:
    binary = format(value, f"0{CHROMOSOME_LENGTH}b")
    return binary_to_gray(binary)


def decode(chromosome: str) -> int:
    return int(gray_to_binary(chromosome), 2)
'''
    else:
        coding_code = '''
def encode(value: int) -> str:
    return format(value, f"0{CHROMOSOME_LENGTH}b")


def decode(chromosome: str) -> int:
    return int(chromosome, 2)
'''

    if ga.parent_select == "best_with_best":
        parent_select_code = '''
def form_parent_pairs(population: List[Individual]) -> List[Tuple[Individual, Individual]]:
    sorted_pop = sorted(population, key=lambda ind: ind.fitness, reverse=True)
    pairs = [(sorted_pop[i], sorted_pop[i + 1]) for i in range(0, len(sorted_pop) - 1, 2)]
    random.shuffle(pairs)
    return pairs
'''
    elif ga.parent_select == "tournament":
        parent_select_code = '''
def tournament_select_one(candidates: List[Individual], tournament_size: int = TOURNAMENT_SIZE) -> Individual:
    tournament = random.sample(candidates, min(tournament_size, len(candidates)))
    return max(tournament, key=lambda ind: ind.fitness)


def form_parent_pairs(population: List[Individual]) -> List[Tuple[Individual, Individual]]:
    return [(tournament_select_one(population), tournament_select_one(population))
            for _ in range(len(population) // 2)]
'''
    else:
        parent_select_code = '''
def form_parent_pairs(population: List[Individual]) -> List[Tuple[Individual, Individual]]:
    shuffled = population[:]
    random.shuffle(shuffled)
    return [(shuffled[i], shuffled[i + 1]) for i in range(0, len(shuffled) - 1, 2)]
'''

    if ga.reduction == "roulette":
        reduction_code = '''
def reduction(population: List[Individual], offspring: List[Individual], new_size: int) -> List[Individual]:
    candidates = population + offspring
    fitness_sum = sum(ind.fitness for ind in candidates)
    if fitness_sum <= 0:
        return random.sample(candidates, new_size)
    probabilities = [ind.fitness / fitness_sum for ind in candidates]
    indices = np.random.choice(len(candidates), size=new_size, replace=True, p=probabilities)
    return [candidates[i] for i in indices]
'''
    else:
        if ga.parent_select == "tournament":
            reduction_code = '''
def reduction(population: List[Individual], offspring: List[Individual], new_size: int) -> List[Individual]:
    candidates = population + offspring
    return [tournament_select_one(candidates) for _ in range(new_size)]
'''
        else:
            reduction_code = '''
def tournament_select_one(candidates: List[Individual], tournament_size: int = TOURNAMENT_SIZE) -> Individual:
    tournament = random.sample(candidates, min(tournament_size, len(candidates)))
    return max(tournament, key=lambda ind: ind.fitness)


def reduction(population: List[Individual], offspring: List[Individual], new_size: int) -> List[Individual]:
    candidates = population + offspring
    return [tournament_select_one(candidates) for _ in range(new_size)]
'''

    template = Template(r'''import math
import random
from dataclasses import dataclass
from typing import List, Tuple

import matplotlib.pyplot as plt
import numpy as np


# ============================================================
# ПАРАМЕТРЫ ЗАДАЧИ. Вариант $variant
# ============================================================

X_MIN = $x_min
X_MAX = $x_max
POPULATION_SIZE = $population_size
CHROMOSOME_LENGTH = $chromosome_length
GENERATIONS = $generations
PC = $pc
PM = $pm
TOURNAMENT_SIZE = $tournament_size
RANDOM_SEED = 42
MAX_INT = 2 ** CHROMOSOME_LENGTH - 1


def target_function(x: float) -> float:
    return $func_code


def calculate_fitness_shift() -> float:
    xs = np.linspace(X_MIN, X_MAX, 5000)
    ys = np.array([target_function(x) for x in xs])
    if "$extreme" == "min":
        return float(np.max(ys))
    return float(np.min(ys))


FITNESS_SHIFT = calculate_fitness_shift()
FITNESS_EPS = 1e-9


def fitness_function(x: float) -> float:
    if "$extreme" == "min":
        return FITNESS_SHIFT - target_function(x) + FITNESS_EPS
    return target_function(x) - FITNESS_SHIFT + FITNESS_EPS


# КОДИРОВАНИЕ И ДЕКОДИРОВАНИЕ ($coding_text)
$coding_code

def int_to_x(value: int) -> float:
    return X_MIN + value * (X_MAX - X_MIN) / MAX_INT


def chromosome_to_x(chromosome: str) -> float:
    return int_to_x(decode(chromosome))


@dataclass
class Individual:
    chromosome: str

    @property
    def x(self) -> float:
        return chromosome_to_x(self.chromosome)

    @property
    def f(self) -> float:
        return target_function(self.x)

    @property
    def fitness(self) -> float:
        return fitness_function(self.x)


def generate_random_individual() -> Individual:
    value = random.randint(0, MAX_INT)
    return Individual(encode(value))


def generate_initial_population(size: int) -> List[Individual]:
    return [generate_random_individual() for _ in range(size)]

$parent_select_code

def crossover_one_point(parent1: Individual, parent2: Individual, pc: float) -> Tuple[Individual, Individual]:
    if random.random() > pc:
        return Individual(parent1.chromosome), Individual(parent2.chromosome)
    point = random.randint(1, CHROMOSOME_LENGTH - 1)
    child1 = parent1.chromosome[:point] + parent2.chromosome[point:]
    child2 = parent2.chromosome[:point] + parent1.chromosome[point:]
    return Individual(child1), Individual(child2)


def point_mutation(individual: Individual, pm: float) -> Individual:
    chromosome = list(individual.chromosome)
    if random.random() < pm:
        bit_index = random.randint(0, CHROMOSOME_LENGTH - 1)
        chromosome[bit_index] = "1" if chromosome[bit_index] == "0" else "0"
    return Individual("".join(chromosome))

$reduction_code

def average_fitness(population: List[Individual]) -> float:
    return sum(ind.fitness for ind in population) / len(population)


def best_individual(population: List[Individual]) -> Individual:
    return max(population, key=lambda ind: ind.fitness)


def genetic_algorithm():
    random.seed(RANDOM_SEED)
    np.random.seed(RANDOM_SEED)
    population = generate_initial_population(POPULATION_SIZE)
    all_populations = [population[:]]
    avg_fitness_history = [average_fitness(population)]

    for _ in range(1, GENERATIONS + 1):
        parent_pairs = form_parent_pairs(population)
        offspring = []
        for parent1, parent2 in parent_pairs:
            child1, child2 = crossover_one_point(parent1, parent2, PC)
            offspring.extend([point_mutation(child1, PM), point_mutation(child2, PM)])
        population = reduction(population, offspring, POPULATION_SIZE)
        all_populations.append(population[:])
        avg_fitness_history.append(average_fitness(population))

    return all_populations, avg_fitness_history


def main():
    print("ГЕНЕТИЧЕСКИЙ АЛГОРИТМ ПОИСКА $extreme_text")
    print("Вариант:", $variant)
    print("Формирование родительских пар: $parent_select_text")
    print("Редукция: $reduction_text")
    all_populations, avg_fitness_history = genetic_algorithm()
    best = best_individual(all_populations[-1])
    print(f"x = {best.x:.10f}")
    print(f"Хромосома = {best.chromosome}")
    print(f"f(x) = {best.f:.10f}")
    print(f"F(x) = {best.fitness:.10f}")

    xs = np.linspace(X_MIN, X_MAX, 2000)
    ys = [target_function(x) for x in xs]
    plt.figure(figsize=(10, 6))
    plt.plot(xs, ys, label="f(x)")
    plt.grid(True)
    plt.legend()
    plt.tight_layout()
    plt.show()


if __name__ == "__main__":
    main()
''')

    return textwrap.dedent(template.substitute(
        variant=ga.variant,
        x_min=ga.x_min,
        x_max=ga.x_max,
        population_size=ga.population_size,
        chromosome_length=ga.chromosome_length,
        generations=ga.generations,
        pc=ga.pc,
        pm=ga.pm,
        tournament_size=ga.tournament_size,
        func_code=func_code,
        extreme=ga.extreme,
        extreme_text=extreme_text,
        coding_text=coding_text,
        parent_select_text=parent_select_text,
        reduction_text=reduction_text,
        coding_code=textwrap.dedent(coding_code).strip(),
        parent_select_code=textwrap.dedent(parent_select_code).strip(),
        reduction_code=textwrap.dedent(reduction_code).strip(),
    ))

def replace_picture_marker(doc: Document, marker: str, image_bytes: bytes, width_inches: float = 5.8):
    """Заменяет абзац с меткой {{graph_...}} на картинку."""
    marker_text = "{{" + marker + "}}"

    def process_paragraph(paragraph):
        if marker_text not in paragraph.text:
            return False
        for run in paragraph.runs:
            run.text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        paragraph.add_run().add_picture(tmp_path, width=Inches(width_inches))
        os.remove(tmp_path)
        return True

    for p in doc.paragraphs:
        if process_paragraph(p):
            return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if process_paragraph(p):
                        return


def fig_to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def make_target_function_fig(ga):
    xs = np.linspace(ga.x_min, ga.x_max, 2000)
    ys = [ga.target_func(x) for x in xs]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(xs, ys, linewidth=1.5)
    ax.set_xlabel("x")
    ax.set_ylabel("f(x)")
    ax.set_title(f"Целевая функция, вариант {ga.variant}")
    ax.grid(True, alpha=0.3)
    return fig


def make_fitness_function_fig(ga):
    xs = np.linspace(ga.x_min, ga.x_max, 2000)
    ys = [ga.fitness(x) for x in xs]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(xs, ys, linewidth=1.5)
    ax.set_xlabel("x")
    ax.set_ylabel("F(x)")
    ax.set_title("Функция приспособленности")
    ax.grid(True, alpha=0.3)
    return fig


def make_population_fig(population, title):
    xs = [ind.x for ind in population]
    ys = [ind.f for ind in population]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.scatter(xs, ys, s=20)
    ax.set_xlabel("x")
    ax.set_ylabel("f(x)")
    ax.set_title(title)
    ax.grid(True, alpha=0.3)
    return fig


def make_avg_fitness_fig(avg_fitness):
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(avg_fitness, marker="o", markersize=3, linewidth=1.5)
    ax.set_xlabel("Поколение")
    ax.set_ylabel("Средняя приспособленность")
    ax.set_title("Динамика средней приспособленности")
    ax.grid(True, alpha=0.3)
    return fig


def function_text_by_variant(variant: int) -> str:
    """Текст функции для отчета. Формулы совпадают с вариантами Task1GA."""
    funcs = {
        1: "2*x - 3.3*x*cos(9.3*x) - 2.3*sin(1.7*x)",
        2: "exp(-0.5*x^2) * sign(cos(9.3*x - 1))",
        3: "-0.6*x + 5.3*|cos(6.1*x)|*cos(3.6*x)",
        4: "-1.3*sin(1.6*x^2 - 0.3)*exp(-0.3*x + 0.5)",
        5: "0.9*|sin(3.7*x)|*cos(6.1*x)",
        6: "0.8*x + 1.4*cos(1.8*x^2)*exp(0.4*x)",
        7: "-sin(0.9*x - 1) - sin(1.8*x - 1)*cos(7.8*x)",
        8: "0.8*x + 1.1*x*sin(9.3*x) - 0.7*cos(0.8*x)",
        9: "0.2*x - 1.1*exp(-0.4*x^2)*sign(cos(9.5*x + 1.5))",
        10: "0.3*x + x*cos(7.3*x) - 0.7*sin(1.3*x)",
        11: "5.1*|cos(15*x)|*cos(33*x)",
        12: "3.5*x + sin(47*x^2 + 2) - 6*x^2",
        13: "9*x - 9.7*|sin(37*x)|*cos(27*x)",
        14: "3.1*x - 3*sin(31.4*x^2 + 7) - 6*x^2",
        15: "6.9*x + 7.5*|cos(39*x)|*cos(31*x)",
        16: "0.7*x + 1.4*x*sin(43*x) - 2.3*cos(5.4*x)",
        17: "cos(1.2*x - 2) - cos(1.7*x - 1)*sin(8.4*x)",
        18: "-0.35*x + 1.2*sin(2.5*x^2)*exp(0.3*x)",
        19: "-0.5*x - 0.7*x*cos(9.2*x + 3) + 0.9*sin(0.9*x - 1)",
        20: "-0.9*|cos(4.1*x)|*sin(6.8*x)",
        21: "-0.5*x + 1.2*exp(-0.5*x^2)*sign(sin(9.7*x + 1.3))",
        22: "-0.9*x + 1.5*cos(2.3*x^2 - 0.5)*exp(-0.4*x + 0.3)",
        23: "0.2*x + x*sin(8.9*x) + 0.9*cos(1.5*x - 2)",
    }
    return funcs.get(variant, "")


def fill_population_values(values: dict, population, prefix: str):
    """Заполняет строки 1-5 и 46-50, как размечено в вашем шаблоне."""
    for row_num in list(range(1, 6)) + list(range(46, 51)):
        ind = population[row_num - 1] if row_num - 1 < len(population) else None
        values[f"{prefix}_x_{row_num}"] = f"{ind.x:.10f}" if ind else ""
        values[f"{prefix}_chromos_{row_num}"] = ind.chromosome if ind else ""
        values[f"{prefix}_f_{row_num}"] = f"{ind.f:.10f}" if ind else ""
        values[f"{prefix}_F_{row_num}"] = f"{ind.fitness:.10f}" if ind else ""


def generate_task1_report_docx(template_bytes: bytes, ga, populations: dict, avg_fitness: list,
                                student_name: str = "", figures: dict = None) -> bytes:
    """Создает docx-отчет по меткам из вашего шаблона."""
    template_bytes = normalize_docx_placeholders(template_bytes)
    doc = Document(io.BytesIO(template_bytes))

    last_generation = max(populations.keys())
    pop0 = populations.get(0, [])
    pop3 = populations.get(3, populations.get(last_generation, []))
    pop_last = populations.get(last_generation, [])
    best_last = max(pop_last, key=lambda ind: ind.fitness)

    values = {
        "name": student_name,
        "func": function_text_by_variant(ga.variant),
        "interval": f"[{ga.x_min}; {ga.x_max}]",
        "extreme": "минимум" if ga.extreme == "min" else "максимум",
        "code": get_program_code_for_report(ga),  # <-- передаем ga
        "p_c": ga.pc,
        "p_m": ga.pm,
        "parent_select": ga.parent_select,
        "reduction": ga.reduction,
        "best_x": f"{best_last.x:.10f}",
        "best_chromosome": best_last.chromosome,
        "best_f": f"{best_last.f:.10f}",
        "best_fitness": f"{best_last.fitness:.10f}",
    }
    fill_population_values(values, pop0, "pop_0")
    fill_population_values(values, pop3, "pop_3")
    fill_population_values(values, pop_last, "pop")

    replace_text_in_doc(doc, values)

    # Используем переданные графики вместо создания новых
    # Если figures не передан или какой-то график отсутствует - создаем заново
    if figures is None:
        figures = {}

    # Графики
    graph_map = {
        "graph_target_func": figures.get('target_func', make_target_function_fig(ga)),
        "graph_fitness": figures.get('fitness', make_fitness_function_fig(ga)),
        "graph_pop_0": figures.get('population_0', make_population_fig(pop0, "Начальная популяция")),
        "graph_pop_3": figures.get('population_3', make_population_fig(pop3, "Третья популяция")),
        "graph_pop": figures.get(f'population_{last_generation}', make_population_fig(pop_last, "Последняя популяция")),
        "graph_avg_fitness": figures.get('avg_fitness', make_avg_fitness_fig(avg_fitness)),
    }

    for marker, fig in graph_map.items():
        replace_picture_marker(doc, marker, fig_to_png_bytes(fig))
        # Закрываем только если это не график из st.session_state (чтобы не закрыть оригинал)
        if figures.get(marker.replace('graph_', '')) is None:
            plt.close(fig)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

# ============================================================
# БЛОК ЗАДАНИЯ №2 (дополнительное) - Поиск бинарных КП
# ============================================================

class Task2GA:
    def __init__(self, variant: int):
        self.variant = variant
        self.set_params_by_variant()

    def set_params_by_variant(self):
        variants = {
            1: {"n": 25, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.85, "pm": 0.15,
                "parent_select": "random", "selection": "roulette"},
            2: {"n": 27, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.9, "pm": 0.15,
                "parent_select": "random", "selection": "tournament"},
            3: {"n": 29, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.8, "pm": 0.1,
                "parent_select": "best_with_best", "selection": "roulette"},
            4: {"n": 31, "population_size": 70, "psl_limit": 2, "k": 4, "pk": 0.85, "pm": 0.15,
                "parent_select": "best_with_best", "selection": "tournament"},
            5: {"n": 33, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.9, "pm": 0.2,
                "parent_select": "random", "selection": "roulette"},
            6: {"n": 35, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.8, "pm": 0.2,
                "parent_select": "best_with_best", "selection": "roulette"},
            7: {"n": 37, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.9, "pm": 0.1,
                "parent_select": "random", "selection": "tournament"},
            8: {"n": 39, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.85, "pm": 0.15,
                "parent_select": "random", "selection": "roulette"},
            9: {"n": 26, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.8, "pm": 0.15,
                "parent_select": "random", "selection": "tournament"},
            10: {"n": 28, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.85, "pm": 0.15,
                 "parent_select": "random", "selection": "roulette"},
            11: {"n": 30, "population_size": 50, "psl_limit": 2, "k": 4, "pk": 0.9, "pm": 0.1,
                 "parent_select": "best_with_best", "selection": "tournament"},
            12: {"n": 32, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.8, "pm": 0.2,
                 "parent_select": "random", "selection": "roulette"},
            13: {"n": 34, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.8, "pm": 0.1,
                 "parent_select": "best_with_best", "selection": "roulette"},
            14: {"n": 36, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.9, "pm": 0.15,
                 "parent_select": "random", "selection": "tournament"},
            15: {"n": 38, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.85, "pm": 0.15,
                 "parent_select": "random", "selection": "roulette"},
            16: {"n": 29, "population_size": 60, "psl_limit": 2, "k": 4, "pk": 0.9, "pm": 0.2,
                 "parent_select": "random", "selection": "tournament"},
            17: {"n": 31, "population_size": 70, "psl_limit": 2, "k": 4, "pk": 0.8, "pm": 0.1,
                 "parent_select": "random", "selection": "roulette"},
            18: {"n": 33, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.8, "pm": 0.15,
                 "parent_select": "best_with_best", "selection": "tournament"},
            19: {"n": 35, "population_size": 70, "psl_limit": 3, "k": 3, "pk": 0.9, "pm": 0.15,
                 "parent_select": "random", "selection": "tournament"},
            20: {"n": 37, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.8, "pm": 0.2,
                 "parent_select": "best_with_best", "selection": "roulette"},
            21: {"n": 39, "population_size": 70, "psl_limit": 4, "k": 3, "pk": 0.9, "pm": 0.2,
                 "parent_select": "best_with_best", "selection": "tournament"},
            22: {"n": 31, "population_size": 100, "psl_limit": 2, "k": 5, "pk": 0.85, "pm": 0.15,
                 "parent_select": "random", "selection": "roulette"},
            23: {"n": 35, "population_size": 100, "psl_limit": 3, "k": 4, "pk": 0.9, "pm": 0.2,
                 "parent_select": "best_with_best", "selection": "tournament"},
        }

        params = variants.get(self.variant)
        if params is None:
            raise ValueError(f"Вариант {self.variant} не найден")

        self.n = params["n"]
        self.population_size = params["population_size"]
        self.psl_limit = params["psl_limit"]
        self.k = params["k"]
        self.pk = params["pk"]
        self.pm = params["pm"]
        self.parent_select = params["parent_select"]
        self.selection = params["selection"]
        self.generations = 200
        self.elite_size = 2

    @dataclass
    class Individual:
        genes: List[int]

        @property
        def sequence_str(self) -> str:
            return "".join("+" if g == 1 else "-" for g in self.genes)

    def aperiodic_acf(self, sequence: List[int]):
        """Вычисление апериодической АКФ"""
        seq = np.array(sequence, dtype=int)
        n = len(seq)

        # Вычисляем АКФ для положительных сдвигов
        acf_positive = []
        for k in range(n):
            if k == 0:
                acf_positive.append(n)
            else:
                val = np.sum(seq[:n - k] * seq[k:])
                acf_positive.append(val)

        # АКФ симметрична
        acf_negative = acf_positive[1:][::-1]
        shifts = list(range(-(n - 1), n))
        values = acf_negative + acf_positive

        return shifts, values

    def positive_sidelobe_level(self, sequence: List[int]) -> int:
        """PSL = максимальный положительный боковой лепесток АКФ"""
        _, acf_vals = self.aperiodic_acf(sequence)
        center = self.n - 1
        sidelobes = [acf_vals[i] for i in range(len(acf_vals)) if i != center]
        positive = [v for v in sidelobes if v > 0]
        return max(positive) if positive else 0

    def fitness(self, sequence: List[int]) -> float:
        """FFn = N / PSL"""
        psl = self.positive_sidelobe_level(sequence)
        if psl == 0:
            return self.n * 10.0
        return self.n / psl

    def generate_random_individual(self):
        genes = [random.choice([-1, 1]) for _ in range(self.n)]
        return self.Individual(genes)

    def generate_initial_population(self):
        return [self.generate_random_individual() for _ in range(self.population_size)]

    def form_parent_pairs(self, population):
        if self.parent_select == "random":
            shuffled = population[:]
            random.shuffle(shuffled)
            pairs = []
            for i in range(0, len(shuffled) - 1, 2):
                pairs.append((shuffled[i], shuffled[i + 1]))
            return pairs
        else:  # best_with_best
            sorted_pop = sorted(population, key=lambda ind: self.fitness(ind.genes), reverse=True)
            pairs = []
            for i in range(0, len(sorted_pop) - 1, 2):
                pairs.append((sorted_pop[i], sorted_pop[i + 1]))
            random.shuffle(pairs)
            return pairs

    def crossover_one_point(self, p1, p2):
        if random.random() > self.pk:
            return self.Individual(p1.genes[:]), self.Individual(p2.genes[:])
        point = random.randint(1, self.n - 1)
        c1 = p1.genes[:point] + p2.genes[point:]
        c2 = p2.genes[:point] + p1.genes[point:]
        return self.Individual(c1), self.Individual(c2)

    def mutation_inversion(self, individual):
        if random.random() > self.pm:
            return individual
        genes = individual.genes[:]
        idx = random.randint(0, self.n - 1)
        genes[idx] *= -1
        return self.Individual(genes)

    def selection_roulette_with_elite(self, candidates):
        sorted_candidates = sorted(candidates, key=lambda ind: self.fitness(ind.genes), reverse=True)
        elite = sorted_candidates[:self.elite_size]

        rest = candidates
        rest_size = self.population_size - self.elite_size

        fitness_vals = [self.fitness(ind.genes) for ind in rest]
        total = sum(fitness_vals)
        if total <= 0:
            probs = [1 / len(rest)] * len(rest)
        else:
            probs = [f / total for f in fitness_vals]

        indices = np.random.choice(len(rest), size=rest_size, replace=True, p=probs)
        selected = [rest[i] for i in indices]

        return elite + selected

    def selection_tournament(self, candidates, tournament_size=3):
        new_pop = []
        candidates_list = list(candidates)
        for _ in range(self.population_size):
            tournament = random.sample(candidates_list, min(tournament_size, len(candidates_list)))
            winner = max(tournament, key=lambda ind: self.fitness(ind.genes))
            new_pop.append(self.Individual(winner.genes[:]))
        return new_pop

    def run(self, seed=42):
        random.seed(seed)
        np.random.seed(seed)

        population = self.generate_initial_population()
        populations = {0: population[:]}
        avg_fitness_history = [sum(self.fitness(ind.genes) for ind in population) / self.population_size]
        best_psl_history = [
            self.positive_sidelobe_level(max(population, key=lambda ind: self.fitness(ind.genes)).genes)]

        found_sequences = {}

        for gen in range(1, self.generations + 1):
            parent_pairs = self.form_parent_pairs(population)
            offspring = []
            for p1, p2 in parent_pairs:
                c1, c2 = self.crossover_one_point(p1, p2)
                c1 = self.mutation_inversion(c1)
                c2 = self.mutation_inversion(c2)
                offspring.extend([c1, c2])

            # Добиваем до нужного размера
            while len(offspring) < self.population_size:
                offspring.append(self.generate_random_individual())

            candidates = population + offspring
            if self.selection == "roulette":
                population = self.selection_roulette_with_elite(candidates)
            else:
                population = self.selection_tournament(candidates)

            if gen == 3:
                populations[3] = population[:]
            if gen == self.generations:
                populations[self.generations] = population[:]

            avg_fitness_history.append(sum(self.fitness(ind.genes) for ind in population) / self.population_size)
            best_ind = max(population, key=lambda ind: self.fitness(ind.genes))
            best_psl_history.append(self.positive_sidelobe_level(best_ind.genes))

            # Сохраняем найденные последовательности
            for ind in population:
                psl = self.positive_sidelobe_level(ind.genes)
                if psl <= self.psl_limit:
                    key = ind.sequence_str
                    if key not in found_sequences:
                        found_sequences[key] = ind

        # Сортируем найденные по качеству
        found_list = list(found_sequences.values())
        found_list.sort(key=lambda ind: self.positive_sidelobe_level(ind.genes))

        return populations, avg_fitness_history, best_psl_history, found_list[:self.k]



# ============================================================
# ОТЧЕТ ДЛЯ ЗАДАНИЯ №2
# ============================================================

def task2_parent_select_text(value: str) -> str:
    return {
        "random": "случайное формирование родительских пар",
        "best_with_best": "формирование пар по принципу «лучшие с лучшими»",
    }.get(value, value)


def task2_selection_text(value: str) -> str:
    return {
        "roulette": "селекция рулеткой с элитизмом",
        "tournament": "турнирная селекция",
    }.get(value, value)


def task2_population_rows(population, ga_obj, limit: int = 15):
    rows = []
    for i, ind in enumerate(population[:limit], 1):
        psl = ga_obj.positive_sidelobe_level(ind.genes)
        ffn = ga_obj.fitness(ind.genes)
        rows.append([str(i), ind.sequence_str, str(psl), f"{ffn:.6f}"])
    return rows


def add_docx_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_docx_key_values(doc: Document, data: List[Tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    for key, value in data:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)
    doc.add_paragraph("")


def add_docx_table(doc: Document, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_values in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_values):
            row[i].text = str(value)
    doc.add_paragraph("")


def add_code_block_to_doc(doc: Document, code: str):
    for line in code.splitlines():
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def make_task2_dynamics_fig(avg_fitness, best_psl_history, ga2):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 4))
    ax1.plot(avg_fitness, linewidth=1.5)
    ax1.set_xlabel("Поколение")
    ax1.set_ylabel("Средняя FFn")
    ax1.set_title("Средняя приспособленность")
    ax1.grid(True, alpha=0.3)

    ax2.plot(best_psl_history, linewidth=1.5)
    ax2.axhline(y=ga2.psl_limit, linestyle="--", label=f"Цель: PSL ≤ {ga2.psl_limit}")
    ax2.set_xlabel("Поколение")
    ax2.set_ylabel("Лучший PSL")
    ax2.set_title("Динамика лучшего PSL")
    ax2.legend()
    ax2.grid(True, alpha=0.3)
    fig.tight_layout()
    return fig


def make_task2_acf_fig(ga2, individual, title: str):
    shifts, acf_vals = ga2.aperiodic_acf(individual.genes)
    psl = ga2.positive_sidelobe_level(individual.genes)
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.plot(shifts, acf_vals, marker="o", markersize=3, linewidth=1.2)
    ax.axhline(y=0, linestyle="-", linewidth=0.5)
    ax.axvline(x=0, linestyle="-", linewidth=0.5)
    center_idx = ga2.n - 1
    for j, (shift, val) in enumerate(zip(shifts, acf_vals)):
        if j != center_idx and val > 0:
            ax.plot(shift, val, "o", markersize=5, zorder=5)
    ax.set_xlabel("Сдвиг k")
    ax.set_ylabel("R(k)")
    ax.set_title(f"{title} (PSL = {psl})")
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    return fig


def get_task2_program_code_for_report(ga2) -> str:
    """Возвращает код решения задания 2 для вставки в отчет."""
    import inspect
    try:
        class_source = inspect.getsource(Task2GA)
    except (OSError, TypeError):
        class_source = "# Не удалось автоматически получить исходный код класса Task2GA."

    header = (
        f"# Параметры варианта {ga2.variant}\n"
        f"# N = {ga2.n}\n"
        f"# P = {ga2.population_size}\n"
        f"# PSL_limit = {ga2.psl_limit}\n"
        f"# K = {ga2.k}\n"
        f"# Pk = {ga2.pk}\n"
        f"# Pm = {ga2.pm}\n"
        f"# Формирование пар: {task2_parent_select_text(ga2.parent_select)}\n"
        f"# Селекция: {task2_selection_text(ga2.selection)}\n"
    )
    launch = (
        "\n\n# Пример запуска:\n"
        "# ga = Task2GA(variant)\n"
        "# populations, avg_fitness, best_psl_history, found_sequences = ga.run(seed=42)\n"
    )
    return header + "\n" + class_source + launch


def generate_task2_report_docx(ga2, populations: dict, avg_fitness: list,
                               best_psl_history: list, found_sequences: list,
                               student_name: str = "") -> bytes:
    """Создает автономный Word-отчет по заданию №2."""
    doc = Document()

    add_docx_heading(doc, "Отчет по заданию №2", level=0)
    if student_name:
        doc.add_paragraph(f"ФИО: {student_name}")
    doc.add_paragraph("Тема: поиск бинарных кодовых последовательностей генетическим алгоритмом.")

    add_docx_heading(doc, "1. Параметры варианта", level=1)
    add_docx_key_values(doc, [
        ("Вариант", ga2.variant),
        ("Длина кодовой последовательности N", ga2.n),
        ("Размер популяции P", ga2.population_size),
        ("Требуемый положительный боковой лепесток PSL", f"≤ {ga2.psl_limit}"),
        ("Количество искомых последовательностей K", ga2.k),
        ("Количество поколений", ga2.generations),
        ("Вероятность скрещивания Pk", ga2.pk),
        ("Вероятность мутации Pm", ga2.pm),
        ("Формирование родительских пар", task2_parent_select_text(ga2.parent_select)),
        ("Селекция", task2_selection_text(ga2.selection)),
        ("Функция приспособленности", "FFn = N / PSL"),
    ])

    add_docx_heading(doc, "2. Динамика работы алгоритма", level=1)
    dynamics_fig = make_task2_dynamics_fig(avg_fitness, best_psl_history, ga2)
    doc.add_picture(io.BytesIO(fig_to_png_bytes(dynamics_fig)), width=Inches(6.3))
    plt.close(dynamics_fig)

    last_generation = max(populations.keys())
    pop0 = populations.get(0, [])
    pop3 = populations.get(3, populations.get(last_generation, []))
    pop_last = populations.get(last_generation, [])

    add_docx_heading(doc, "3. Популяции", level=1)
    for title, population in [
        ("Начальная популяция", pop0),
        ("Третья популяция", pop3),
        ("Последняя популяция", pop_last),
    ]:
        add_docx_heading(doc, title, level=2)
        add_docx_table(doc, ["№", "КП", "PSL", "FFn"], task2_population_rows(population, ga2, limit=15))

    add_docx_heading(doc, "4. Найденные кодовые последовательности", level=1)
    if found_sequences:
        found_rows = []
        for i, ind in enumerate(found_sequences[:ga2.k], 1):
            psl = ga2.positive_sidelobe_level(ind.genes)
            ffn = ga2.fitness(ind.genes)
            found_rows.append([str(i), ind.sequence_str, str(psl), f"{ffn:.6f}"])
        add_docx_table(doc, ["№", "КП", "PSL", "FFn"], found_rows)

        for i, ind in enumerate(found_sequences[:ga2.k], 1):
            add_docx_heading(doc, f"АКФ найденной КП №{i}", level=2)
            acf_fig = make_task2_acf_fig(ga2, ind, f"АКФ найденной КП №{i}")
            doc.add_picture(io.BytesIO(fig_to_png_bytes(acf_fig)), width=Inches(6.3))
            plt.close(acf_fig)
    else:
        doc.add_paragraph(f"За {ga2.generations} поколений не найдено КП с PSL ≤ {ga2.psl_limit}.")
        if pop_last:
            best_ind = max(pop_last, key=lambda ind: ga2.fitness(ind.genes))
            best_psl = ga2.positive_sidelobe_level(best_ind.genes)
            best_ffn = ga2.fitness(best_ind.genes)
            add_docx_key_values(doc, [
                ("Лучшая найденная КП", best_ind.sequence_str),
                ("Лучший PSL", best_psl),
                ("FFn", f"{best_ffn:.6f}"),
            ])
            acf_fig = make_task2_acf_fig(ga2, best_ind, "АКФ лучшей найденной КП")
            doc.add_picture(io.BytesIO(fig_to_png_bytes(acf_fig)), width=Inches(6.3))
            plt.close(acf_fig)

    add_docx_heading(doc, "5. Программный код", level=1)
    add_code_block_to_doc(doc, get_task2_program_code_for_report(ga2))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ============================================================
# STREAMLIT UI
# ============================================================
def run():
    st.title("🧬 Генетические алгоритмы")
    st.markdown("Автоматизированное решение заданий по генетическим алгоритмам")

    task = st.sidebar.radio("Выберите задание", ["Задание №1 (Поиск экстремума)", "Задание №2 (Поиск КП)"])

    if task == "Задание №1 (Поиск экстремума)":
        st.header("📈 Задание №1: Поиск экстремума целевой функции")

        variant_options = list(range(1, 24))
        variant = st.selectbox("Выберите вариант (1-23)", variant_options, format_func=lambda x: f"Вариант {x}")

        generations = st.slider("Количество поколений", 10, 100, 30)
        seed = st.number_input("Seed (для воспроизводимости)", value=42, step=1)

        if st.button("🚀 Запустить ГА", type="primary"):
            with st.spinner("Выполняется генетический алгоритм..."):
                ga = Task1GA(variant)
                ga.generations = generations
                populations, avg_fitness = ga.run(seed=seed)
                st.session_state.task1_results = {
                    'ga': ga,
                    'populations': populations,
                    'avg_fitness': avg_fitness,
                    'variant': variant,
                    'generations': generations,
                    'seed': seed
                }

            st.success("Выполнено!")
            st.session_state.task1_figures = {}

            # ===== ВЫВОД НАЧАЛЬНЫХ ДАННЫХ =====
            st.subheader("📋 Начальные данные")

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Вариант", variant)
            with col2:
                st.metric("Область поиска", f"[{ga.x_min}; {ga.x_max}]")
            with col3:
                st.metric("Экстремум", "Минимум" if ga.extreme == "min" else "Максимум")
            with col4:
                st.metric("Кодирование", "Двоичный код" if ga.coding == "binary" else "Код Грея")

            col5, col6, col7, col8 = st.columns(4)
            with col5:
                st.metric("Размер популяции", ga.population_size)
            with col6:
                st.metric("Длина хромосомы", f"{ga.chromosome_length} бит")
            with col7:
                st.metric("Pc (скрещивание)", f"{ga.pc}")
            with col8:
                st.metric("Pm (мутация)", f"{ga.pm}")

            st.markdown(f"**Формирование родительских пар:** {ga.parent_select}")
            st.markdown(f"**Редукция:** {ga.reduction}")

            # ===== ГРАФИК ЦЕЛЕВОЙ ФУНКЦИИ =====
            st.subheader("📈 График целевой функции")
            xs = np.linspace(ga.x_min, ga.x_max, 2000)
            ys = [ga.target_func(x) for x in xs]

            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(xs, ys, label='f(x)', linewidth=1.5)
            ax.set_xlabel('x')
            ax.set_ylabel('f(x)')
            ax.set_title(f'Целевая функция (Вариант {variant})')
            ax.grid(True, alpha=0.3)
            st.pyplot(fig)
            st.session_state.task1_figures['target_func'] = fig

            # ===== ГРАФИК ФУНКЦИИ ПРИСПОСОБЛЕННОСТИ =====
            st.subheader("📊 График функции приспособленности")
            fitness_ys = [ga.fitness(x) for x in xs]

            fig2, ax2 = plt.subplots(figsize=(10, 5))
            ax2.plot(xs, fitness_ys, color='green', linewidth=1.5)
            ax2.set_xlabel('x')
            ax2.set_ylabel('F(x)')
            ax2.set_title('Функция приспособленности')
            ax2.grid(True, alpha=0.3)
            st.pyplot(fig2)
            st.session_state.task1_figures['fitness'] = fig2

            # ===== ГРАФИКИ ПОПУЛЯЦИЙ =====
            st.subheader("👥 Визуализация популяций")

            pop_indices = [0, 3, generations] if generations >= 3 else [0, generations]

            for pop_idx in pop_indices:
                if pop_idx in populations:
                    population = populations[pop_idx]

                    fig3, ax3 = plt.subplots(figsize=(10, 5))
                    ax3.plot(xs, ys, label='f(x)', linewidth=1.5, alpha=0.7)

                    pop_x = [ind.x for ind in population]
                    pop_f = [ind.f for ind in population]
                    ax3.scatter(pop_x, pop_f, c='red', s=40, alpha=0.7, label='Особи популяции', zorder=5)

                    ax3.set_xlabel('x')
                    ax3.set_ylabel('f(x)')
                    ax3.set_title(f'Популяция на графике целевой функции (поколение {pop_idx})')
                    ax3.legend()
                    ax3.grid(True, alpha=0.3)
                    st.pyplot(fig3)
                    st.session_state.task1_figures[f'population_{pop_idx}'] = fig3

            # ===== ТАБЛИЦЫ ПОПУЛЯЦИЙ =====
            st.subheader("📋 Таблицы популяций")


            def display_population_table(population, ga_obj, title, pop_num):
                """Вывод таблицы: первые 10 и последние 10 особей"""
                st.markdown(f"**{title} (поколение {pop_num})**")

                # Первые 10
                st.caption("📌 Первые 10 особей")
                data_first = []
                for i, ind in enumerate(population[:10], 1):
                    data_first.append({
                        "№": i,
                        "x": f"{ind.x:.8f}",
                        "Хромосома": ind.chromosome,
                        "f(x)": f"{ind.f:.8f}",
                        "F(x)": f"{ind.fitness:.8f}"
                    })
                st.dataframe(data_first, use_container_width=True)

                # Последние 10
                if len(population) > 10:
                    st.caption("📌 Последние 10 особей")
                    data_last = []
                    start_idx = len(population) - 10
                    for i, ind in enumerate(population[start_idx:], start_idx + 1):
                        data_last.append({
                            "№": i,
                            "x": f"{ind.x:.8f}",
                            "Хромосома": ind.chromosome,
                            "f(x)": f"{ind.f:.8f}",
                            "F(x)": f"{ind.fitness:.8f}"
                        })
                    st.dataframe(data_last, use_container_width=True)

                avg_f = sum(ind.fitness for ind in population) / len(population)
                st.caption(f"📊 Средняя приспособленность: {avg_f:.8f}")
                st.markdown("---")


            # Выводим таблицы для 0, 3 и последней популяции
            if 0 in populations:
                display_population_table(populations[0], ga, "НАЧАЛЬНАЯ ПОПУЛЯЦИЯ", 0)

            if 3 in populations:
                display_population_table(populations[3], ga, "ТРЕТЬЯ ПОПУЛЯЦИЯ", 3)
            elif generations < 3 and generations in populations:
                display_population_table(populations[generations], ga, f"ПОПУЛЯЦИЯ (поколение {generations})", generations)

            if generations in populations and generations != 0 and generations != 3:
                display_population_table(populations[generations], ga, "ПОСЛЕДНЯЯ ПОПУЛЯЦИЯ", generations)

            # ===== ГРАФИК СРЕДНЕЙ ПРИСПОСОБЛЕННОСТИ =====
            st.subheader("📈 Изменение средней приспособленности")
            fig4, ax4 = plt.subplots(figsize=(10, 5))
            ax4.plot(avg_fitness, marker='o', markersize=4, linewidth=1.5)
            ax4.set_xlabel('Поколение')
            ax4.set_ylabel('Средняя приспособленность')
            ax4.set_title('Динамика средней приспособленности')
            ax4.grid(True, alpha=0.3)
            st.pyplot(fig4)
            st.session_state.task1_figures['avg_fitness'] = fig4

            # ===== ИТОГОВЫЙ РЕЗУЛЬТАТ =====
            st.subheader("🏆 Итоговый результат")

            best_last = max(populations[generations], key=lambda ind: ind.fitness)

            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Лучший x", f"{best_last.x:.10f}")
            with col2:
                st.metric("Хромосома", best_last.chromosome)
            with col3:
                st.metric("f(x)", f"{best_last.f:.10f}")
            with col4:
                st.metric("F(x)", f"{best_last.fitness:.10f}")


            st.markdown("---")
            st.subheader("Сформировать отчет Word")

            student_name = st.text_input("ФИО для титульного листа", value="")

            template_path = Path(__file__).parent / "patter.docx"

            if not template_path.exists():
                st.error(f"Шаблон отчета не найден: {template_path.name}")
            else:
                # Передаем сохраненные графики в функцию отчета
                report_bytes = generate_task1_report_docx(
                    template_path.read_bytes(),
                    ga,
                    populations,
                    avg_fitness,
                    student_name=student_name,
                    figures=st.session_state.task1_figures  # <-- передаем сохраненные графики
                )

                st.download_button(
                    "Скачать отчет .docx",
                    data=report_bytes,
                    file_name=f"report_task1_variant_{variant}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


    else:  # Задание 2
        st.header("🔢 Задание №2: Поиск бинарных кодовых последовательностей")

        variant_options = list(range(1, 24))
        variant = st.selectbox("Выберите вариант (1-23)", variant_options, format_func=lambda x: f"Вариант {x}")

        generations = st.slider("Количество поколений", 50, 500, 200, key="gen2")
        seed = st.number_input("Seed (для воспроизводимости)", value=42, step=1, key="seed2")

        if st.button("🚀 Запустить ГА (поиск КП)", type="primary"):
            with st.spinner("Выполняется генетический алгоритм поиска КП..."):
                ga2 = Task2GA(variant)
                ga2.generations = generations
                populations, avg_fitness, best_psl_history, found_seqs = ga2.run(seed=seed)

            st.success("Выполнено!")
            st.session_state.task2_results = {
                'ga': ga2,
                'populations': populations,
                'avg_fitness': avg_fitness,
                'best_psl_history': best_psl_history,
                'found_seqs': found_seqs,
                'variant': variant,
                'generations': generations,
                'seed': seed,
            }
            st.session_state.task2_figures = {}

            # ===== ПАРАМЕТРЫ ВАРИАНТА =====
            st.subheader("📋 Параметры варианта")

            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.metric("N (длина КП)", ga2.n)
            with col2:
                st.metric("Размер популяции (P)", ga2.population_size)
            with col3:
                st.metric("Требуемый PSL", ga2.psl_limit)
            with col4:
                st.metric("Искомое K", ga2.k)
            with col5:
                st.metric("Pk / Pm", f"{ga2.pk} / {ga2.pm}")

            col6, col7 = st.columns(2)
            with col6:
                st.metric("Формирование пар", "Случайный" if ga2.parent_select == "random" else "Лучшие с лучшими")
            with col7:
                st.metric("Селекция", "Рулетка" if ga2.selection == "roulette" else "Турнир")

            # ===== ГРАФИКИ =====
            st.subheader("📊 Динамика работы алгоритма")

            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

            # График средней приспособленности
            ax1.plot(avg_fitness, linewidth=1.5, color='blue')
            ax1.set_xlabel('Поколение')
            ax1.set_ylabel('Средняя FFn')
            ax1.set_title('Средняя приспособленность')
            ax1.grid(True, alpha=0.3)

            # График лучшего PSL
            ax2.plot(best_psl_history, linewidth=1.5, color='red')
            ax2.axhline(y=ga2.psl_limit, color='green', linestyle='--', label=f'Цель: PSL ≤ {ga2.psl_limit}')
            ax2.set_xlabel('Поколение')
            ax2.set_ylabel('Лучший PSL')
            ax2.set_title('Динамика лучшего PSL')
            ax2.legend()
            ax2.grid(True, alpha=0.3)

            plt.tight_layout()
            st.pyplot(fig)

            # ===== ТАБЛИЦЫ ПОПУЛЯЦИЙ =====
            st.subheader("📋 Популяции")


            def display_kp_table(population, ga_obj, title, pop_num):
                st.markdown(f"**{title} (поколение {pop_num})**")

                data = []
                for i, ind in enumerate(population[:15], 1):
                    psl = ga_obj.positive_sidelobe_level(ind.genes)
                    ffn = ga_obj.fitness(ind.genes)
                    data.append({
                        "№": i,
                        "КП": ind.sequence_str,
                        "PSL": psl,
                        "FFn": f"{ffn:.4f}"
                    })
                st.dataframe(data, use_container_width=True)

                avg_psl = sum(ga_obj.positive_sidelobe_level(ind.genes) for ind in population) / len(population)
                avg_ffn = sum(ga_obj.fitness(ind.genes) for ind in population) / len(population)

                col1, col2 = st.columns(2)
                with col1:
                    st.caption(f"📊 Средний PSL: {avg_psl:.2f}")
                with col2:
                    st.caption(f"📊 Средняя FFn: {avg_ffn:.4f}")
                st.markdown("---")


            # Выводим таблицы
            if 0 in populations:
                display_kp_table(populations[0], ga2, "НАЧАЛЬНАЯ ПОПУЛЯЦИЯ", 0)

            if 3 in populations:
                display_kp_table(populations[3], ga2, "ТРЕТЬЯ ПОПУЛЯЦИЯ", 3)

            if generations in populations:
                display_kp_table(populations[generations], ga2, "ПОСЛЕДНЯЯ ПОПУЛЯЦИЯ", generations)

            # ===== ГРАФИКИ АКФ ДЛЯ НАЙДЕННЫХ КП =====
            st.subheader("🔍 Найденные кодовые последовательности")

            if found_seqs:
                st.success(f"Найдено {len(found_seqs)} КП с PSL ≤ {ga2.psl_limit}")

                for i, seq in enumerate(found_seqs[:ga2.k]):
                    psl = ga2.positive_sidelobe_level(seq.genes)
                    ffn = ga2.fitness(seq.genes)

                    with st.expander(f"КП #{i + 1} | PSL = {psl}, FFn = {ffn:.4f}", expanded=True):
                        st.markdown(f"**Последовательность:** `{seq.sequence_str}`")
                        st.markdown(f"**Длина N = {ga2.n}**")
                        st.markdown(f"**PSL = {psl}** (требуемый ≤ {ga2.psl_limit})")

                        # График АКФ
                        shifts, acf_vals = ga2.aperiodic_acf(seq.genes)

                        fig_acf, ax_acf = plt.subplots(figsize=(12, 5))

                        # Строим АКФ
                        ax_acf.plot(shifts, acf_vals, marker='o', markersize=4, linewidth=1.5, color='purple')
                        ax_acf.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
                        ax_acf.axvline(x=0, color='black', linestyle='-', linewidth=0.5)

                        # Выделяем боковые лепестки
                        center_idx = ga2.n - 1
                        for j, (shift, val) in enumerate(zip(shifts, acf_vals)):
                            if j != center_idx and val > 0:
                                ax_acf.plot(shift, val, 'ro', markersize=6, zorder=5)

                        ax_acf.set_xlabel('Сдвиг k')
                        ax_acf.set_ylabel('R(k)')
                        ax_acf.set_title(f'АКФ найденной КП #{i + 1} (PSL = {psl})')
                        ax_acf.grid(True, alpha=0.3)

                        st.pyplot(fig_acf)
            else:
                st.warning(f"⚠️ За {generations} поколений не найдено КП с PSL ≤ {ga2.psl_limit}")

                # Показываем лучшую найденную
                if generations in populations:
                    best_ind = max(populations[generations], key=lambda ind: ga2.fitness(ind.genes))
                    best_psl = ga2.positive_sidelobe_level(best_ind.genes)
                    st.info(f"📌 Лучший достигнутый PSL: {best_psl} (при требуемом {ga2.psl_limit})")

                    # График АКФ лучшей
                    shifts, acf_vals = ga2.aperiodic_acf(best_ind.genes)
                    fig_acf, ax_acf = plt.subplots(figsize=(12, 5))
                    ax_acf.plot(shifts, acf_vals, marker='o', markersize=4, linewidth=1.5)
                    ax_acf.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
                    ax_acf.axvline(x=0, color='black', linestyle='-', linewidth=0.5)
                    ax_acf.set_xlabel('Сдвиг k')
                    ax_acf.set_ylabel('R(k)')
                    ax_acf.set_title(f'АКФ лучшей найденной КП (PSL = {best_psl})')
                    ax_acf.grid(True, alpha=0.3)
                    st.pyplot(fig_acf)

                st.markdown("**Рекомендации для улучшения:**")
                st.markdown("- Увеличьте количество поколений (рекомендуется 300-500)")
                st.markdown("- Увеличьте размер популяции")
                st.markdown("- Измените seed для другого случайного поиска")
                st.markdown("- Попробуйте другой вариант")

        if st.session_state.task2_results is not None:
            result = st.session_state.task2_results
            ga2_report = result['ga']
            st.markdown("---")
            st.subheader("Сформировать отчет Word по заданию №2")
            student_name_2 = st.text_input("ФИО для титульного листа", value="", key="student_name_task2")

            report_bytes_2 = generate_task2_report_docx(
                ga2_report,
                result['populations'],
                result['avg_fitness'],
                result['best_psl_history'],
                result['found_seqs'],
                student_name=student_name_2,
            )

            st.download_button(
                "Скачать отчет по заданию №2 .docx",
                data=report_bytes_2,
                file_name=f"report_task2_variant_{result['variant']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_task2_report",
            )
